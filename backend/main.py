from bs4 import BeautifulSoup
import requests
from duckduckgo_search import DDGS
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.responses import FileResponse, PlainTextResponse
from pydantic import BaseModel
import asyncio
import uvicorn
import os
import re
import difflib
import csv
import shutil
import docx
import json
import logging
import hashlib
import sys
import time
from typing import Optional, List
from datetime import datetime
from dotenv import load_dotenv
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

from fastapi.responses import RedirectResponse
from services.ai_service import analyze_resume, generate_cover_letter, analyze_job_metadata, generate_additional_points, generate_recruiter_outreach_email, generate_checkin_followup_email, generate_linkedin_message, extract_contacts_from_text
from services.ollama_service import generate_mail_draft, generate_follow_up, detect_w2_fulltime
from services.docx_service import extract_text_from_docx, create_tailored_docx, insert_bullets_after
from services import gmail_service
from services.profile_service import process_uploaded_doc
from services.usage_tracker import get_usage_stats
from database import init_db, save_resume_record, save_job_matcher_record, get_all_resumes, delete_resume_record, update_resume_status, get_resume_by_id, find_existing_company, sanitize_csv_field, search_records, update_user_address, save_follow_up_draft, get_follow_up_draft, update_resume_after_edit, update_user_notes, update_resume_rerun

DATA_DIR = os.getenv("DATA_DIR", "data")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(os.path.join(DATA_DIR, "logs"), exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(DATA_DIR, 'logs', 'app.log')),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Validate required environment variables
load_dotenv()
required_env_vars = ["GEMINI_API_KEY"]
for var in required_env_vars:
    if not os.getenv(var):
        logger.error(f"Missing required environment variable: {var}")
        raise SystemExit(f"Missing required environment variable: {var}")

limiter = Limiter(key_func=get_remote_address)

_TECH_TOOLS = {
    "digitalocean", "digital ocean", "aws", "amazon", "amazon web services",
    "google cloud", "gcp", "azure", "microsoft azure", "kubernetes", "docker",
    "terraform", "ansible", "jenkins", "github", "gitlab", "bitbucket",
    "circleci", "travis", "datadog", "splunk", "grafana", "prometheus",
    "elasticsearch", "redis", "mongodb", "postgresql", "mysql", "nginx",
    "apache", "linux", "ubuntu", "centos", "helm", "argocd", "vault",
    "consul", "nomad", "istio", "envoy", "kafka", "rabbitmq", "jira",
    "confluence", "slack", "pagerduty", "new relic", "cloudflare",
    "heroku", "vercel", "netlify", "firebase", "supabase",
}

_JOB_TITLE_WORDS = {
    "senior", "junior", "lead", "staff", "principal", "manager", "director",
    "engineer", "developer", "architect", "analyst", "administrator", "admin",
    "devops", "sre", "cloud", "software", "platform", "infrastructure",
    "backend", "frontend", "full stack", "fullstack", "remote", "hybrid",
    "onsite", "contract", "full time", "part time", "position", "role",
    "job description", "job title", "responsibilities", "requirements",
    "qualifications", "experience", "we are looking", "about the role",
}

_GENERIC_EMAIL = {"gmail", "yahoo", "hotmail", "outlook", "protonmail",
                  "mail", "icloud", "aol", "live", "ymail", "zoho",
                  "fastmail", "tutanota", "pm", "hey"}

RESUME_FILENAME = "Teja_Mahesh_Neerukonda_Resume.docx"

def _is_valid_company_name(name: str) -> bool:
    """Filter out names that are job titles, tech tools, or generic words."""
    lower = name.lower().strip()
    if len(lower) < 3 or len(lower) > 60:
        return False
    if lower in _TECH_TOOLS:
        return False
    words = set(lower.split())
    if words & _JOB_TITLE_WORDS and len(words - _JOB_TITLE_WORDS) == 0:
        return False
    return True

def _scrape_jd_from_url(url: str) -> str:
    """Fetch a URL and extract the job description using structured data + AI cleanup."""
    import requests
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self._texts = []
            self._skip = False
            self._skip_tags = {'script', 'style', 'nav', 'footer', 'noscript', 'svg', 'head'}
            self._json_ld = []
            self._in_json_ld = False
            self._current_tag = None

        def handle_starttag(self, tag, attrs):
            self._current_tag = tag
            if tag in self._skip_tags:
                self._skip = True
            attrs_dict = dict(attrs)
            if tag == 'script' and attrs_dict.get('type') == 'application/ld+json':
                self._in_json_ld = True
                self._skip = False

        def handle_endtag(self, tag):
            if tag in self._skip_tags:
                self._skip = False
            if tag == 'script':
                self._in_json_ld = False

        def handle_data(self, data):
            if self._in_json_ld:
                self._json_ld.append(data)
            elif not self._skip:
                text = data.strip()
                if text:
                    self._texts.append(text)

        def get_text(self):
            return '\n'.join(self._texts)

        def get_json_ld(self):
            return ''.join(self._json_ld)

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.9',
    }
    MAX_CONTENT_BYTES = 3 * 1024 * 1024  # 3 MB cap to prevent memory spike
    resp = requests.get(url, headers=headers, timeout=15, stream=True)
    resp.raise_for_status()
    content_chunks = []
    total = 0
    for chunk in resp.iter_content(chunk_size=65536, decode_unicode=False):
        total += len(chunk)
        content_chunks.append(chunk)
        if total >= MAX_CONTENT_BYTES:
            break
    resp.close()
    raw_bytes = b''.join(content_chunks)
    encoding = resp.encoding or 'utf-8'
    html_text = raw_bytes.decode(encoding, errors='replace')

    parser = _TextExtractor()
    parser.feed(html_text)

    # Try JSON-LD structured data first (LinkedIn, Indeed, etc. embed this)
    jd_from_structured = ""
    json_ld_raw = parser.get_json_ld()
    if json_ld_raw:
        try:
            ld_data = json.loads(json_ld_raw)
            if isinstance(ld_data, list):
                ld_data = next((d for d in ld_data if d.get('@type') == 'JobPosting'), ld_data[0] if ld_data else {})
            if isinstance(ld_data, dict):
                parts = []
                if ld_data.get('title'):
                    parts.append(f"Job Title: {ld_data['title']}")
                org = ld_data.get('hiringOrganization', {})
                if isinstance(org, dict) and org.get('name'):
                    parts.append(f"Company: {org['name']}")
                elif isinstance(org, str):
                    parts.append(f"Company: {org}")
                loc = ld_data.get('jobLocation', {})
                if isinstance(loc, dict):
                    addr = loc.get('address', {})
                    if isinstance(addr, dict):
                        loc_parts = [addr.get('addressLocality', ''), addr.get('addressRegion', ''), addr.get('addressCountry', '')]
                        loc_str = ', '.join(p for p in loc_parts if p)
                        if loc_str:
                            parts.append(f"Location: {loc_str}")
                if ld_data.get('employmentType'):
                    parts.append(f"Employment Type: {ld_data['employmentType']}")
                desc = ld_data.get('description', '')
                if desc:
                    # Strip HTML tags from description
                    import re as _re
                    desc_clean = _re.sub(r'<[^>]+>', ' ', desc)
                    desc_clean = _re.sub(r'\s+', ' ', desc_clean).strip()
                    parts.append(f"\n{desc_clean}")
                if parts:
                    jd_from_structured = '\n'.join(parts)
        except (json.JSONDecodeError, StopIteration, TypeError):
            pass

    if jd_from_structured and len(jd_from_structured) >= 100:
        return jd_from_structured

    # Fallback: use raw text but ask AI to extract just the JD
    raw_text = parser.get_text()
    if len(raw_text) < 50:
        raise ValueError("Could not extract enough text from the URL")

    # Use AI to clean up the scraped text
    try:
        from services.ai_service import client as ai_client
        from services.model_config import GEMINI_QUALITY_MODEL
        import google.genai as genai
        from google.genai import types

        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            client = genai.Client(api_key=api_key)
            response = client.models.generate_content(
                model=GEMINI_QUALITY_MODEL,
                contents=f"""Extract ONLY the job description from this scraped web page text.
Remove all navigation, ads, cookie notices, sidebar content, and other non-JD text.
Return the clean job description including: job title, company name, location, requirements, responsibilities, qualifications, and any other JD-related content.
If you cannot find a job description, return "NO_JD_FOUND".

Scraped text (first 8000 chars):
{raw_text[:8000]}""",
                config=types.GenerateContentConfig(temperature=0.1),
            )
            cleaned = response.text.strip()
            if cleaned and cleaned != "NO_JD_FOUND" and len(cleaned) >= 50:
                from services.usage_tracker import log_api_call
                usage = response.usage_metadata
                if usage:
                    log_api_call(GEMINI_QUALITY_MODEL, "url_jd_extract",
                                 input_tokens=usage.prompt_token_count or 0,
                                 output_tokens=(usage.candidates_token_count or 0) + (usage.thoughts_token_count or 0))
                return cleaned
    except Exception as e:
        logger.warning(f"AI JD extraction fallback failed: {e}")

    return raw_text

def _extract_company_name(jd_text: str) -> str:
    """Best-effort local extraction of company name — no API call.
    Searches the ENTIRE JD text including signatures at the bottom."""

    # Priority 1: Look for explicit company name lines (often near signatures)
    company_patterns = [
        r'(?:Company|Employer|Organization|Firm|Client|Hiring\s+Company)(?:\s+Name)?[:\s]+([^\n]{2,60})',
        r'^([A-Z][A-Za-z0-9\s&.,\-]+?)\s+(?:is hiring|is looking|is seeking|seeks|are hiring|has an opening)',
        r'(?:About|Join|Work\s+at|Working\s+at)\s+([A-Z][A-Za-z0-9\s&.,\-]{2,40}?)(?:\s*[\n:]|\s+is\b|\s+are\b)',
        r'(?:at|@)\s+([A-Z][A-Za-z0-9\s&.\-]{2,40}?)(?:\s*[\n,.]|\s+in\b|\s+for\b|\s*$)',
        r'(?:Role|Position|Title|Job|Opportunity)\s*(?:at|with|@|[-–])\s*([A-Z][A-Za-z0-9\s&.\-]{2,40})',
        r'^([A-Z][A-Za-z0-9\s&.\-]{2,40}?)\s*[-–|]\s*(?:Job|Role|Position|Career|Hiring|Opening)',
        r'(?:Employer|Hiring\s+Company)[:\s]+([^\n]{2,60})',
        r'(?:on behalf of|client|end[\s-]?client)[:\s]+([^\n]{2,60})',
    ]
    for pattern in company_patterns:
        m = re.search(pattern, jd_text, re.MULTILINE | re.IGNORECASE)
        if m:
            name = m.group(1).strip().strip('.,')
            if _is_valid_company_name(name):
                return name

    # Priority 2: Extract from "Company Inc/LLC/Ltd/Corp" patterns anywhere in text
    for line in jd_text.split('\n'):
        corp_match = re.search(
            r'([A-Z][A-Za-z0-9]+(?:[\s&]+[A-Z][A-Za-z0-9]+)*)\s+(?:Inc\.?|LLC|Ltd\.?|Corp\.?|Corporation|Pvt\.?|Private|Limited|Group|Consulting|Solutions|Technologies|Services|Staffing|Partners|Associates)',
            line.strip()
        )
        if corp_match:
            name = corp_match.group(0).strip().strip('.,')
            if _is_valid_company_name(name):
                return name

    # Priority 3: Extract from email domain (search ALL emails in the JD)
    all_emails = re.findall(r'[\w.+-]+@([\w-]+)\.\w{2,}', jd_text)
    for domain in all_emails:
        if domain.lower() not in _GENERIC_EMAIL and len(domain) > 2:
            return domain.replace('-', ' ').title()

    # Priority 4: Look for "Web: www.company.com" or website patterns near signatures
    web_match = re.search(r'(?:Web|Website|URL|Visit\s+us)[:\s]+(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+)\.\w{2,}', jd_text, re.IGNORECASE)
    if web_match:
        domain = web_match.group(1).strip()
        if domain.lower() not in _GENERIC_EMAIL and len(domain) > 2:
            return domain.replace('-', ' ').title()

    # Priority 5: Any URL in the text (e.g. https://company.com/careers)
    _URL_BLACKLIST = {'linkedin', 'indeed', 'glassdoor', 'ziprecruiter', 'lever',
                      'greenhouse', 'workday', 'icims', 'smartrecruiters', 'jobvite',
                      'google', 'apply', 'dice', 'monster', 'careerbuilder', 'hired',
                      'angel', 'wellfound', 'simplyhired', 'recruiter', 'zoom', 'teams'}
    url_match = re.search(r'https?://(?:www\.)?([a-zA-Z0-9-]+)\.(?:com|io|org|co|net|ai)\b', jd_text, re.IGNORECASE)
    if url_match:
        domain = url_match.group(1).strip()
        if domain.lower() not in _GENERIC_EMAIL and domain.lower() not in _URL_BLACKLIST and len(domain) > 2:
            return domain.replace('-', ' ').title()

    # Priority 6: Phone number with company signature — look for a name on the line before a phone
    lines = jd_text.strip().split('\n')
    for i, line in enumerate(lines):
        if re.search(r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}', line):
            if i > 0:
                prev = lines[i - 1].strip()
                if prev and len(prev) < 50 and not re.match(r'^[\d(+]', prev):
                    candidate = re.sub(r'[,|–\-].*$', '', prev).strip()
                    if _is_valid_company_name(candidate) and not re.search(r'@', candidate):
                        return candidate

    return "Unknown_Company"


def _guess_title_from_job_text(jt: str) -> str:
    """Best-effort job title for a raw scraped/aggregated job text blob, used to label
    postings that get filtered out by hard-reject rules before reaching AI scoring."""
    marker = 'Title:'
    if jt.startswith(marker):
        rest = jt[len(marker):]
        for sep in ('\\nCompany:', '\nCompany:', '\\n', '\n'):
            if sep in rest:
                return rest.split(sep)[0].strip()[:80]
        return rest.strip()[:80]
    return jt.strip().split('\n')[0][:80].strip()


def _make_company_dir(company_name: str, email: str = "", root: str = "trailerd") -> str:
    """Build a unique <root>/<company> directory path. Uses email if unknown, appends _1, _2 if duplicate.
    root must be one of the two allowed storage roots — 'trailerd' for the Dashboard's regular
    resume tailoring, 'online-platform' for jobs tailored from the Command Center."""
    if root not in ("trailerd", "online-platform"):
        root = "trailerd"
    safe = "".join([c for c in company_name if c.isalpha() or c.isdigit() or c == ' ' or c == '@' or c == '.']).strip()
    # Windows silently strips trailing dots/spaces from a path component when looking it
    # up (os.path.exists/isdir/open all miss it) even though it creates the folder with
    # the literal name — a company like "ATSIT Inc." would create a folder Windows can
    # never find again. Strip them here so the name we create is the name we can read.
    safe = safe.rstrip('. ')
    if not safe or safe.lower() in ("unknown company", "unknowncompany", "unknown", "unknown_company"):
        if email and isinstance(email, str) and '@' in email:
            safe = email.split('@')[0]
            safe = "".join([c for c in safe if c.isalnum()]).strip()

        if not safe or safe.lower() in ("unknown company", "unknowncompany", "unknown", "unknown_company"):
            safe = "Unknown"

    safe = safe.replace(' ', '_')
    base_dir = f"{root}/{safe}"
    
    # Handle duplicates by appending _1, _2, etc. — check-then-create (os.path.exists
    # followed by a separate os.makedirs) has a TOCTOU race: two concurrent requests
    # for the same company (e.g. a double-submit) can both see the directory as "not
    # existing" and both settle on the same path, then both write into it and silently
    # clobber each other's output. Making the directory itself the atomic check (via
    # exist_ok=False raising FileExistsError) closes that window.
    company_dir = base_dir
    counter = 1
    while True:
        try:
            os.makedirs(company_dir, exist_ok=False)
            break
        except FileExistsError:
            company_dir = f"{base_dir}_{counter}"
            counter += 1
    return company_dir


def _resolve_resume_path(record: dict) -> Optional[str]:
    """Best-effort locate a history record's tailored resume file on disk. The stored
    file_path can go stale even though the company's own folder still exists — e.g. a
    legacy scan stored a flat "<root>/<Company>.docx" path from before per-company
    subfolders existed, or the file inside the folder was renamed/regenerated by hand.
    Falls back to searching that folder (and a freshly-derived company-name folder)
    before giving up, and self-heals the DB record so future lookups skip the fallback."""
    file_path = record.get('file_path')
    if file_path and os.path.exists(file_path):
        return file_path

    # Bare storage roots hold every company's data side by side — never treat one as
    # "this record's folder" or a stray file for a different company could get grabbed.
    bare_roots = {os.path.abspath("trailerd"), os.path.abspath("online-platform")}

    search_dirs = []
    company_name = record.get('company_name')
    # A generic placeholder name ("Unknown", "Unknown Company", ...) is shared by many
    # unrelated records — matching on it would silently attach a stranger's resume to
    # this record, so only trust a company-name-derived folder for a real company name.
    if company_name and company_name.strip().lower().replace('_', ' ') not in (
        "unknown company", "unknown", ""
    ):
        safe = "".join(c for c in company_name if c.isalnum() or c in (' ', '.')).strip().rstrip('. ').replace(' ', '_')
        if safe:
            search_dirs.append(f"trailerd/{safe}")
            search_dirs.append(f"online-platform/{safe}")

    if file_path:
        search_dirs.append(os.path.dirname(file_path))
        # Legacy flat-file scans stored "<root>/<Company>.docx" with no subfolder —
        # a same-named folder for that company may still exist one level "down".
        stem, ext = os.path.splitext(file_path)
        if ext:
            search_dirs.append(stem)

    seen = set()
    for d in search_dirs:
        if not d or d in seen or not os.path.isdir(d) or os.path.abspath(d) in bare_roots:
            continue
        seen.add(d)
        docs = sorted(
            f for f in os.listdir(d)
            if f.endswith(".docx") and "cover" not in f.lower() and not f.startswith("cover_letter_")
        )
        if docs:
            found = os.path.join(d, docs[0]).replace("\\", "/")
            if found != file_path and record.get('id'):
                try:
                    from database import link_tailored_resume
                    link_tailored_resume(record['id'], found)
                except Exception as e:
                    logger.warning(f"Could not self-heal file_path for record {record.get('id')}: {e}")
            return found
    return None


def _convert_docx_to_pdf(docx_path: str) -> Optional[str]:
    """Best-effort docx -> pdf conversion via MS Word COM automation (docx2pdf) — never
    raises. Requires MS Word installed on the host, so it silently no-ops wherever that
    isn't true (e.g. the Linux/Alpine Docker image); a scan/tailoring action must still
    succeed even without a PDF copy."""
    try:
        from docx2pdf import convert
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        convert(docx_path, pdf_path)
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        logger.warning(f"docx-to-pdf conversion failed for {docx_path}: {e}")
        return None


def _job_artifact_dir(record_id: int, company_name: str) -> str:
    """Deterministic online-platform/<company>_<job_id> folder for every AI-generated
    artifact tied to one Command Center job — keyed by job id (not name-dedup counters)
    so every action (tailor, cover letter, drafts, contact info, match explanation,
    status changes) always lands in the exact same folder for that job, regardless of
    which action ran first or how the AI spelled the company name."""
    safe = "".join([c for c in (company_name or "") if c.isalpha() or c.isdigit() or c == ' ' or c == '@' or c == '.']).strip()
    safe = safe.replace(' ', '_') or "Unknown"
    job_dir = f"online-platform/{safe}_{record_id}"
    os.makedirs(job_dir, exist_ok=True)
    return job_dir


def _write_job_artifact(record_id: int, company_name: str, filename: str, content: str):
    """Best-effort text artifact write for a Command Center job action. Never raises —
    a disk hiccup here shouldn't break the actual action (draft generation, status
    update, etc.) it's just leaving a record of."""
    try:
        job_dir = _job_artifact_dir(record_id, company_name)
        with open(f"{job_dir}/{filename}", "w", encoding="utf-8") as f:
            f.write(content)
    except Exception as e:
        logger.warning(f"Failed to write job artifact {filename} for job {record_id}: {e}")


def get_output_filename(filename_type: str, record_id: int = 0) -> str:
    """Generate non-revealing output filenames using record ID hash."""
    if record_id > 0:
        user_hash = hashlib.sha256(str(record_id).encode()).hexdigest()[:8]
    else:
        user_hash = hashlib.sha256(str(datetime.now().timestamp()).encode()).hexdigest()[:8]

    templates = {
        "resume": f"resume_{user_hash}.docx",
        "cover_letter": f"cover_letter_{user_hash}.docx",
        "mail_draft": f"mail_draft_{user_hash}.txt"
    }
    return templates.get(filename_type, f"file_{user_hash}.tmp")

class StatusUpdate(BaseModel):
    status: str

class BatchScanRequest(BaseModel):
    jd_texts: List[str]
    selected_resume: str
    ai_notes: Optional[str] = None

def _extract_experience_years(jd_text: str) -> int:
    """Extract the maximum years of experience required from a JD.
    Handles: '11+ Years', '10-15 years experience', '84 months', etc."""
    text_lower = jd_text.lower()
    max_years = 0

    # Range patterns: "10-15 years experience" — this function extracts the MAXIMUM
    # years required, so the capture group must be around the upper bound (15), not the
    # lower bound (10). Capturing the lower bound made a "5-15 years" JD register as
    # only 5 years, silently letting postings requiring far more than the max-10-years
    # cutoff pass the hard-reject filter.
    range_pattern = r'\d{1,2}\s*[\-–to]+\s*(\d{1,2})\s*(?:years?|yrs?)(?:\s+of)?(?:\s+\w+){0,3}\s*(?:experience|exp\b)'
    for m in re.finditer(range_pattern, text_lower):
        years = int(m.group(1))
        if 1 <= years <= 30:
            max_years = max(max_years, years)

    text_no_ranges = re.sub(r'\d{1,2}\s*[\-–to]+\s*\d{1,2}\s*(?:years?|yrs?)', '', text_lower)

    patterns = [
        # "11+ years of experience", "5 years experience"
        r'(\d{1,2})\s*\+?\s*(?:years?|yrs?)(?:\s+of)?(?:\s+\w+){0,3}\s*(?:experience|exp\b|professional)',
        # "minimum 10 years", "at least 8 years", "requires 12 years"
        r'(?:minimum|min\.?|at\s+least|over|requires?|must\s+have)\s*(\d{1,2})\s*\+?\s*(?:years?|yrs?)',
        # "experience: 10+ years"
        r'(?:experience|exp)\s*(?:required)?[:\s]+(\d{1,2})\s*\+?\s*(?:years?|yrs?)',
        # "10+ years in DevOps/cloud/IT"
        r'(\d{1,2})\s*\+?\s*(?:years?|yrs?)\s+(?:in\s+(?:the\s+)?(?:industry|field|role|domain|it\b|devops|cloud|software))',
        # "10+ years working/hands-on/relevant"
        r'(\d{1,2})\s*\+?\s*(?:years?|yrs?)\s+(?:working|hands[\-\s]on|relevant|total|overall|progressive)',
        # "11+ Years Candidates Only" — number + years + any word (catch-all for title lines)
        r'(\d{1,2})\s*\+\s*(?:years?|yrs?)\s+(?:candidates?|only|required|mandatory)',
        # Standalone "11+ years" or "11+ yrs" (with the + sign = explicit requirement)
        r'\b(\d{1,2})\s*\+\s*(?:years?|yrs?)\b',
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text_no_ranges):
            years = int(m.group(1))
            if 1 <= years <= 30:
                max_years = max(max_years, years)

    # Convert months to years: "84 months of experience" = 7 years
    month_patterns = [
        r'(\d{2,3})\s*(?:months?|mos?\.?)(?:\s+of)?(?:\s+\w+){0,3}\s*(?:experience|exp\b)',
        r'(\d{2,3})\s*\+?\s*(?:months?|mos?\.?)\s+(?:in\s+|of\s+)',
    ]
    for pattern in month_patterns:
        for m in re.finditer(pattern, text_lower):
            months = int(m.group(1))
            if 6 <= months <= 360:
                years = months // 12
                max_years = max(max_years, years)

    return max_years

def _check_lead_role(jd_text: str) -> str | None:
    """Return a skip reason if the JD is for a lead/management position, else None."""
    text_lower = jd_text.lower()

    # Check the job title area (first 500 chars where the title usually is)
    title_area = text_lower[:500]

    # Lead patterns
    lead_patterns = [
        r'\blead\s+(?:sre|devops|site\s+reliability)\b',
        r'\blead\s+(?:software|platform|cloud|infrastructure|systems?|backend|frontend|full[\s-]?stack)\s+engineer\b',
        r'\blead\s+(?:developer|architect|admin|administrator|engineer)\b',
        r'\b(?:sre|devops|software|platform|cloud|infrastructure|engineering)\s+lead\b',
        r'\btechnical\s+lead\b',
        r'\btech\s+lead\b',
        r'\bteam\s+lead\b',
        r'\bsr\.?\s+lead\b',
    ]

    # Management/senior leadership patterns (check title area only)
    mgmt_patterns = [
        r'\bengineering\s+manager\b',
        r'\bdevops\s+manager\b',
        r'\binfrastructure\s+manager\b',
        r'\bplatform\s+manager\b',
        r'\bcloud\s+manager\b',
        r'\bsre\s+manager\b',
        r'\bit\s+manager\b',
        r'\bmanager[,\s]+(?:devops|sre|infrastructure|platform|cloud|engineering)\b',
        r'\bdirector\s+of\s+(?:engineering|devops|infrastructure|platform|sre|it|operations|technology)\b',
        r'\bvp\s+(?:of\s+)?(?:engineering|infrastructure|devops|platform|operations|technology)\b',
        r'\bvice\s+president\b',
        r'\bhead\s+of\s+(?:engineering|devops|infrastructure|platform|sre|it|operations)\b',
        r'\bchief\s+(?:technology|information|infrastructure)\s+officer\b',
        r'\b(?:cto|cio)\b',
        r'\bprincipal\s+(?:engineer|architect|devops|sre)\b',
        r'\bstaff\s+(?:engineer|architect|devops|sre)\b',
    ]

    # Both pattern groups are scoped to title_area, not the full JD body — a normal
    # individual-contributor posting can easily mention "reports to the Team Lead" or
    # "collaborates with our DevOps Lead" in its body without the ADVERTISED ROLE
    # itself being a lead position. Searching the whole text (the original bug here)
    # mass-false-rejected legitimate IC postings over an incidental mention elsewhere.
    for pattern in lead_patterns:
        if re.search(pattern, title_area):
            return "JD is for a lead-level position — skipping"

    for pattern in mgmt_patterns:
        if re.search(pattern, title_area):
            return "JD is for a management/senior leadership role — skipping"

    return None

def _check_visa_eligibility(jd_text: str) -> str | None:
    """Return a detailed skip reason if the JD is not GC-friendly, else None."""
    text_lower = jd_text.lower()

    explicit_exclusions = [
        r'\bno\s+green\s*card\b',
        r'\bno\s+gc\b',
        r'\bfake\s+gc\b',
        r'\bfake\s+green\s*card\b',
        r'\bgc\s+not\s+accepted\b',
        r'\bgreen\s*card\s+not\s+accepted\b',
        # "US Citizens Only" unambiguously excludes Green Card holders — this belongs
        # here as an always-reject signal, not in the soft found_visas/has_gc check
        # below, since it's never legitimately "GC-friendly-if-we-just-find-the-word-
        # green-card-elsewhere."
        r'\bus\s+citizens?\s+only\b',
    ]
    for pattern in explicit_exclusions:
        if re.search(pattern, text_lower):
            return "Rejected: JD explicitly excludes Green Card holders"

    gc_patterns = [
        r'\bgreen\s*card\b', r'\bgc\b', r'\bpermanent\s+resident\b',
        r'\bpermanent\s+residency\b',
    ]

    # Only SPECIFIC named visa types are a genuine "this posting is scoped to a
    # particular visa category" signal. Bare "sponsorship"/"visa sponsor" mentions were
    # removed — that's near-universal boilerplate ("no sponsorship available/required")
    # that says nothing about GC eligibility (a GC holder never needs sponsorship
    # either way), and combined with dropping "citizen" as a GC-friendly signal (see
    # git history), it was rejecting nearly every real posting overnight. "W2 Only" was
    # also removed — that's an employment-type signal already enforced by
    # _check_c2c_c2h_only, not a visa-eligibility one, and doesn't belong here.
    visa_labels = {
        r'\bh[- ]?1b\b': 'H-1B',
        r'\bh[- ]?1\b': 'H-1',
        r'\bopt\b': 'OPT',
        r'\bcpt\b': 'CPT',
        r'\btn\s+visa\b': 'TN Visa',
        r'\bl[- ]?1\b': 'L-1',
        r'\bl[- ]?2\b': 'L-2',
        r'\be[- ]?2\b': 'E-2',
        r'\be[- ]?3\b': 'E-3',
        r'\bo[- ]?1\b': 'O-1',
        r'\bf[- ]?1\b': 'F-1',
        r'\bj[- ]?1\b': 'J-1',
        r'\bead\b': 'EAD',
        r'\bwork\s+permit\b': 'Work Permit',
    }

    has_gc = any(re.search(p, text_lower) for p in gc_patterns)
    found_visas = [label for pattern, label in visa_labels.items() if re.search(pattern, text_lower)]

    if found_visas and not has_gc:
        visa_list = ', '.join(sorted(set(found_visas)))
        return f"Rejected: JD requires [{visa_list}] — Green Card not listed as accepted"

    return None

def _check_foreign_language(jd_text: str) -> str | None:
    """Return a skip reason if the JD strictly requires a foreign language, else None."""
    text_lower = jd_text.lower()
    languages = r'(chinese|mandarin|spanish|mexican|french|german|japanese|korean|russian|arabic|portuguese|italian|hindi)'
    
    patterns = [
        languages + r'\s*\(?required\)?',
        languages + r'\s*is\s*required',
        r'\b(?:required|mandatory|must)\b.{0,30}' + languages,
        r'\b(?:fluent|native|bilingual)\b.{0,30}' + languages + r'.{0,30}\b(?:required|mandatory|must)\b',
        languages + r'.{0,30}\b(?:fluent|native|bilingual)\b.{0,30}\b(?:required|mandatory|must)\b'
    ]
    
    for p in patterns:
        if re.search(p, text_lower):
            return "JD explicitly requires a foreign language"
            
    return None

def _check_employment_type(jd_text: str) -> tuple[str, str | None]:
    """
    Check employment type: C2C, C2H, Full-time, W2
    Return: (employment_type, warning_if_rejected)
    User accepts ONLY C2C or C2H contracts, rejects full-time/W2
    """
    text_lower = jd_text.lower()

    # Check for explicit C2C patterns (strong signals only)
    c2c_patterns = [
        r'\bc2c\b', r'\bcorp[\s-]to[\s-]corp\b', r'\bcorp-to-corp\b',
        r'\b1099\b', r'\bindependent\s+contractor\b',
    ]

    # Check for C2H patterns
    c2h_patterns = [
        r'\bc2h\b', r'\bcontract[\s-]to[\s-]hire\b', r'\bcontract[\s-]hire\b',
        r'\bcontract\s+to\s+perm\b',
    ]

    # Check for W2/Full-time patterns
    w2_patterns = [
        r'\bw2\b', r'\bw-2\b',
        r'\bfull[\s-]?time\s+(?:position|role|opportunity|employment|employee)\b',
        r'\bfulltime\b',
        r'\bpermanent\s+(?:position|role|employee|employment)\b',
        r'\bdirect\s+hire\b',
        r'\bsalaried\s+(?:position|role|employee)\b',
    ]

    # Check for contract (weaker signal — only if no W2 signals)
    contract_patterns = [
        r'\bcontract\s+(?:position|role|opportunity|assignment|engagement)\b',
        r'\bcontractor\s+(?:position|role)\b',
        r'\b(?:6|12|18)\s*[\+]?\s*month\s+contract\b',
        r'\bcontract\s+duration\b',
    ]

    has_c2c = any(re.search(p, text_lower) for p in c2c_patterns)
    has_c2h = any(re.search(p, text_lower) for p in c2h_patterns)
    has_w2 = any(re.search(p, text_lower) for p in w2_patterns)
    has_contract = any(re.search(p, text_lower) for p in contract_patterns)

    # W2 check first — if it says "full-time" or "W2", reject even if "contract" appears elsewhere
    if has_w2 and not has_c2c and not has_c2h:
        return ('w2', "This is a W2/Full-time position. You prefer C2C/C2H contracts.")

    if has_c2c:
        return ('c2c', None)
    if has_c2h:
        return ('c2h', None)
    if has_contract:
        return ('contract', None)

    # Default: unclear employment type — pass as soft warning, not hard reject
    return ('unknown', None)


def _check_position_filled(jd_text: str) -> str | None:
    """Return a skip reason if the scraped page shows the posting is already closed/
    filled — this is common on the DDG-scraping fallback path, where the fetched page
    is the live posting and may already say so, even though it still ranked in search."""
    text_lower = jd_text.lower()
    patterns = [
        r'\bposition\s+has\s+been\s+filled\b',
        r'\bjob\s+has\s+been\s+filled\b',
        r'\brole\s+has\s+been\s+filled\b',
        r'\bthis\s+job\s+is\s+no\s+longer\s+available\b',
        r'\bno\s+longer\s+accepting\s+applications\b',
        r'\bno\s+longer\s+accepting\s+applicants\b',
        r'\bposting\s+has\s+expired\b',
        r'\bthis\s+posting\s+has\s+expired\b',
        r'\bjob\s+posting\s+has\s+expired\b',
        r'\bposition\s+(?:is\s+)?closed\b',
        r'\bthis\s+position\s+is\s+no\s+longer\s+open\b',
        r'\bapplications?\s+(?:are\s+)?closed\b',
        r'\bhas\s+been\s+filled\b',
    ]
    for p in patterns:
        if re.search(p, text_lower):
            return "Posting appears to be closed or already filled — skipping before AI scoring"
    return None


def _check_remote_only(jd_text: str) -> str | None:
    """Command Center only wants FULLY remote roles in the US — reject postings that
    are explicitly hybrid or require on-site/in-office work. Hybrid is checked first
    and rejected unconditionally, even if the posting also mentions 'remote' somewhere
    (e.g. "Hybrid — 3 days/week in office, remote flexibility considered" is still not
    fully remote). Absence of the word 'remote' entirely is NOT enough to reject on its
    own (many postings just don't mention work mode) — only an explicit hybrid/onsite
    signal is.

    IMPORTANT: hybrid detection must require WORK-MODE context (e.g. "hybrid role",
    "hybrid work model", "Hybrid (3 days/week)"), never a bare "\\bhybrid\\b" match — a
    huge fraction of DevOps/Cloud postings say things like "hybrid cloud architecture"
    or "hybrid infrastructure", which is a technology description, not a work-location
    one. Matching on the bare word silently mass-rejected genuinely fully-remote
    postings that happened to mention hybrid cloud tech (caught live: a posting titled
    "Senior DevOps Engineer - Remote" was rejected solely because its JD mentioned
    "hybrid cloud")."""
    text_lower = jd_text.lower()

    hybrid_patterns = [
        r'\bhybrid\s+(?:role|position|schedule|arrangement|setup|work(?:place)?)\b',
        r'\bwork\s+model\s*:?\s*hybrid\b',
        r'\bhybrid\s+work\s+model\b',
        r'\b\d+\s*days?\s*(?:/|a|per)?\s*week\s+(?:in\s+(?:the\s+)?office\s+)?hybrid\b',
        r'\bhybrid\s*[:\-]\s*\d+\s*days?\b',
        r'\bhybrid\s*\(\s*\d+\s*days?',
        r'\bhybrid\s+\d+\s*days?\s*(?:a|per)?\s*week\b',
        r'\bhybrid\s*/\s*(?:onsite|remote|on-?site)\b',
        r'\b(?:onsite|remote|on-?site)\s*/\s*hybrid\b',
        r'\(\s*hybrid\s*/',
        r'\b(?:location|work\s+type|work\s+arrangement|worksite)\s*:?\s*hybrid\b',
    ]
    for p in hybrid_patterns:
        if re.search(p, text_lower):
            return "This role is hybrid — you're only searching for fully remote."

    if re.search(r'\bremote\b', text_lower):
        return None

    onsite_patterns = [
        r'\bon-?site\s+(?:only|required|position|role)\b',
        r'\bmust\s+(?:work|be)\s+on-?site\b',
        r'\bin-?office\s+(?:only|required|position|role|5\s*days)\b',
        r'\bno\s+remote\b',
        r'\bremote\s+work\s+is\s+not\s+(?:available|permitted|offered)\b',
        r'\brelocation\s+required\b',
        r'\bmust\s+relocate\b',
        r'\bcandidates?\s+must\s+reside\s+in\b',
        r'\b100%\s+on-?site\b',
        r'\b5\s+days\s+(?:a\s+week\s+)?in\s+(?:the\s+)?office\b',
    ]
    for p in onsite_patterns:
        if re.search(p, text_lower):
            return "This role requires on-site work — you're only searching for remote."
    return None


def _check_experience_years(jd_text: str, max_years: int = 10) -> str | None:
    """Reject postings that require more experience than the candidate has — same
    max-10-years cutoff already enforced on the Dashboard/Job Matcher scan paths
    (_extract_experience_years / max allowed: 10), just not previously wired into
    Command Center's own hard-reject chain."""
    years = _extract_experience_years(jd_text)
    if years > max_years:
        return f"Requires {years}+ years experience (max allowed: {max_years})"
    return None


def _check_c2c_c2h_only(jd_text: str, contract_types: list = None) -> str | None:
    """Reject postings whose employment type isn't one the user actually selected in
    Command Center's search filters (contract_types: any of 'w2'/'c2c'/'c2h' from the
    UI checkboxes). Previously this was hardcoded to a strict C2C/C2H-only allowlist
    REGARDLESS of what the frontend's W2 checkbox was set to — meaning toggling W2 on
    in the UI did nothing, and (combined with the fact that most LinkedIn/Indeed/etc.
    postings are W2 or don't state a contract type at all) silently rejected almost
    every posting, even with W2 selected. Falls back to the old strict C2C/C2H-only
    allowlist only if contract_types is empty/not provided, so any other caller that
    doesn't pass it keeps prior behavior.

    Unspecified/ambiguous employment type (most postings don't explicitly say "C2C" or
    "W2") is treated as compatible with a W2 selection — most unlabeled postings are
    effectively standard W2 hires — but still rejected when only C2C/C2H were
    selected, preserving a genuine C2C/C2H-only search when that's what's wanted."""
    allowed = {t.lower() for t in (contract_types or [])} or {'c2c', 'c2h'}
    allowed_label = '/'.join(sorted(t.upper() for t in allowed))

    employment_type, _ = _check_employment_type(jd_text)

    if employment_type in ('c2c', 'c2h'):
        if employment_type in allowed:
            return None
        return f"This is a {employment_type.upper()} contract position — you're only searching for {allowed_label}."

    if employment_type == 'w2':
        if 'w2' in allowed:
            return None
        return f"This is a W2/Full-time position — you're only searching for {allowed_label}."

    if 'w2' in allowed:
        return None
    return f"No explicit {allowed_label} contract terms found — only searching for {allowed_label}."


def extract_job_keywords(jd_text: str) -> dict:
    """
    Phase 1: Extract tech keywords from job description
    Returns dict with keyword categories and found keywords
    """
    text_lower = jd_text.lower()

    keyword_groups = {
        'devops': [
            'kubernetes', 'k8s', 'docker', 'container', 'orchestration',
            'terraform', 'ansible', 'infrastructure', 'iac', 'infrastructure as code',
            'aws', 'gcp', 'azure', 'cloud', 'ec2', 's3', 'lambda',
            'ci/cd', 'jenkins', 'gitlab', 'github', 'circleci', 'travis',
            'monitoring', 'prometheus', 'datadog', 'grafana', 'elastic',
            'linux', 'bash', 'shell', 'scripting', 'python',
            'helm', 'helm charts', 'argocd', 'flux', 'gitops',
            'nginx', 'apache', 'haproxy', 'load balancer',
        ],
        'sre': [
            'site reliability', 'sre', 'reliability', 'uptime', 'availability',
            'slo', 'sla', 'oncall', 'on-call', 'pagerduty',
            'incident', 'incident response', 'incident management',
            'postmortem', 'blameless', 'runbook',
            'observability', 'logging', 'tracing', 'metrics',
        ],
        'platform': [
            'platform engineer', 'platform team', 'internal platform',
            'developer platform', 'developer experience', 'devex',
            'self-service', 'automation', 'tooling',
        ],
        'security': [
            'security', 'infosec', 'secops', 'infrastructure security',
            'compliance', 'audit', 'vulnerability', 'scanning',
            'iam', 'identity', 'access management', 'iam',
            'encryption', 'tls', 'ssl', 'certificates',
        ],
    }

    found = {}
    for category, keywords in keyword_groups.items():
        found[category] = []
        for keyword in keywords:
            if keyword in text_lower:
                found[category].append(keyword)

    return found

def calculate_match_score(years_required: int, keywords: dict, employment_type: str) -> int:
    """
    Calculate overall job match percentage (0-100)
    Based on: experience years, keyword overlap, employment type match
    """
    score = 100

    # Experience deduction: each year over 10 costs 2 points
    if years_required > 10:
        score -= min(20, (years_required - 10) * 2)

    # Employment type deduction
    if employment_type == 'w2':
        score -= 30  # Major deduction for W2 (not acceptable)
    elif employment_type == 'unknown':
        score -= 10  # Minor deduction for unclear type

    # Keyword match bonus (small positive impact)
    total_keywords = sum(len(kws) for kws in keywords.values())
    if total_keywords >= 3:
        score += 5
    elif total_keywords == 0:
        score -= 10

    return max(0, min(100, score))  # Clamp to 0-100

def _csv_header():
    return ["Timestamp", "Company Name", "Before Score", "After Score", "Status", "Skip Reason", "Original Resume", "Tailored Resume", "JD Info", "Cover Letter", "Mail Draft", "Job Description", "Source", "Source URL", "Job Fit %", "Rejection Reason"]

def append_to_csv(company_name, jd_text, before_score, after_score, original_path, tailored_path):
    csv_file = os.path.join(DATA_DIR, "history.csv")
    file_exists = os.path.exists(csv_file)

    with open(csv_file, mode='a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(_csv_header())

        abs_orig = os.path.abspath(original_path) if original_path else ""
        abs_tail = os.path.abspath(tailored_path) if tailored_path else ""

        tail_dir = os.path.dirname(abs_tail) if abs_tail else ""
        abs_jd = os.path.join(tail_dir, "jd_info.txt") if tail_dir else ""
        abs_cl_search = [f for f in os.listdir(tail_dir) if f.startswith("cover_letter_")] if tail_dir and os.path.isdir(tail_dir) else []
        abs_mail_search = [f for f in os.listdir(tail_dir) if f.startswith("mail_draft_")] if tail_dir and os.path.isdir(tail_dir) else []

        abs_cl = os.path.join(tail_dir, abs_cl_search[0]) if abs_cl_search else ""
        abs_mail = os.path.join(tail_dir, abs_mail_search[0]) if abs_mail_search else ""

        orig_link = f'=HYPERLINK("file:///{abs_orig.replace(chr(92), "/")}", "Open Original")' if abs_orig else "N/A"
        tail_link = f'=HYPERLINK("file:///{abs_tail.replace(chr(92), "/")}", "Open Tailored")' if abs_tail else "N/A"
        jd_link = f'=HYPERLINK("file:///{abs_jd.replace(chr(92), "/")}", "Open JD Info")' if abs_jd else "N/A"
        cl_link = f'=HYPERLINK("file:///{abs_cl.replace(chr(92), "/")}", "Open Cover Letter")' if abs_cl else "N/A"
        mail_link = f'=HYPERLINK("file:///{abs_mail.replace(chr(92), "/")}", "Open Mail Draft")' if abs_mail else "N/A"

        writer.writerow([
            sanitize_csv_field(datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            sanitize_csv_field(company_name),
            sanitize_csv_field(f"{before_score}%"),
            sanitize_csv_field(f"{after_score}%"),
            "Processed",
            "",
            orig_link,
            tail_link,
            jd_link,
            cl_link,
            mail_link,
            sanitize_csv_field(jd_text[:1000]),
            "dashboard",
            "",
            "",
            ""
        ])

def append_skipped_to_csv(jd_text, skip_reason):
    csv_file = os.path.join(DATA_DIR, "history.csv")
    file_exists = os.path.exists(csv_file)

    with open(csv_file, mode='a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(_csv_header())

        company_name = _extract_company_name(jd_text) if len(jd_text) >= 50 else "Unknown"
        writer.writerow([
            sanitize_csv_field(datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            sanitize_csv_field(company_name),
            "N/A",
            "N/A",
            "Skipped",
            sanitize_csv_field(skip_reason),
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            sanitize_csv_field(jd_text[:1000]),
            "dashboard",
            "",
            "",
            ""
        ])

def append_job_matcher_to_csv(company_name, jd_text, match_pct, status, rejection_reason="", source_url=""):
    csv_file = os.path.join(DATA_DIR, "history.csv")
    file_exists = os.path.exists(csv_file)

    with open(csv_file, mode='a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(_csv_header())

        writer.writerow([
            sanitize_csv_field(datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            sanitize_csv_field(company_name),
            "N/A",
            "N/A",
            sanitize_csv_field(status),
            "",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            sanitize_csv_field(jd_text[:1000]),
            "job-finder",
            sanitize_csv_field(source_url or ""),
            sanitize_csv_field(f"{match_pct}%") if match_pct is not None else "N/A",
            sanitize_csv_field(rejection_reason),
        ])

# Ensure output directories exist
os.makedirs("trailerd", exist_ok=True)
os.makedirs("online-platform", exist_ok=True)
os.makedirs("original", exist_ok=True)
init_db()

from prometheus_fastapi_instrumentator import Instrumentator

app = FastAPI(title="Job Tailored Resume API")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Setup Prometheus metrics
Instrumentator().instrument(app).expose(app)

# Security: Trusted host middleware
allowed_hosts = os.getenv("ALLOWED_HOSTS", "localhost,127.0.0.1").split(",")
app.add_middleware(TrustedHostMiddleware, allowed_hosts=[h.strip() for h in allowed_hosts])

# Security: CORS with restricted origins
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[o.strip() for o in allowed_origins],
    allow_credentials=True,
    allow_methods=["GET", "POST", "DELETE", "PATCH"],
    allow_headers=["Content-Type"],
    max_age=3600
)

# Security: Add security headers
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    response.headers["X-Permitted-Cross-Domain-Policies"] = "none"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    return response

@app.get("/")
def read_root():
    return {"status": "ok", "message": "API is running"}

@app.get("/api/base-resume")
def get_base_resume():
    base_path = "original/base_resume.docx"
    if os.path.exists(base_path):
        return {"exists": True, "filename": "base_resume.docx"}
    return {"exists": False}

@app.get("/api/resumes")
@limiter.limit("60/minute")
async def list_resumes(request: Request):
    try:
        resumes = []
        if os.path.exists("original"):
            for f in os.listdir("original"):
                if f.endswith(".docx"):
                    path = os.path.join("original", f)
                    stat = os.stat(path)
                    resumes.append({
                        "filename": f,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                    })
        resumes.sort(key=lambda r: r["modified"], reverse=True)
        return resumes
    except Exception as e:
        logger.error(f"List resumes error: {e}")
        raise HTTPException(status_code=500, detail="Failed to list resumes")

@app.post("/api/resumes")
@limiter.limit("10/minute")
async def upload_named_resume(request: Request, resume: UploadFile = File(...)):
    try:
        if not resume.filename or not resume.filename.endswith(".docx"):
            raise HTTPException(status_code=400, detail="Only .docx files are supported")

        file_bytes = await resume.read()
        if len(file_bytes) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")

        if len(file_bytes) == 0:
            raise HTTPException(status_code=400, detail="File is empty")

        # Sanitize filename - only allow alphanumeric, dash, underscore, dot
        safe_name = re.sub(r"[^\w\-.]", "_", os.path.basename(resume.filename))
        safe_name = re.sub(r"\.{2,}", ".", safe_name)  # Prevent directory traversal

        file_path = os.path.join("original", safe_name)

        # Verify path is within original directory
        abs_path = os.path.abspath(file_path)
        abs_base = os.path.abspath("original")
        if not abs_path.startswith(abs_base):
            raise HTTPException(status_code=400, detail="Invalid filename")

        with open(file_path, "wb") as f:
            f.write(file_bytes)

        logger.info(f"Resume uploaded: {safe_name}")
        return {"filename": safe_name}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume upload error: {e}")
        raise HTTPException(status_code=500, detail="Failed to upload resume")

@app.delete("/api/resumes/{filename}")
@limiter.limit("10/minute")
async def delete_named_resume(request: Request, filename: str):
    try:
        safe = os.path.basename(filename)
        file_path = os.path.join("original", safe)

        # Verify path is within original directory
        abs_path = os.path.abspath(file_path)
        abs_base = os.path.abspath("original")
        if not abs_path.startswith(abs_base):
            raise HTTPException(status_code=404, detail="Resume not found")

        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Resume not found")

        os.remove(file_path)
        logger.info(f"Resume deleted: {safe}")
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume deletion error: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete resume")

async def _scan_resume_core(
    request: Request,
    jd_text: str,
    resume: Optional[UploadFile] = None,
    selected_resume: Optional[str] = None,
    ai_notes: Optional[str] = None,
    rerun_id: Optional[int] = None,
    storage_root: str = "trailerd",
    override_company_dir: Optional[str] = None,
    skip_duplicate_check: bool = False,
):
    """Internal scan/tailor implementation. storage_root, override_company_dir, and
    skip_duplicate_check are internal-only — never accept these from an HTTP request.
    The public /api/scan endpoint below only forwards the plain user-facing fields;
    only trusted internal callers (Command Center tailoring) pass the rest, and even
    then override_company_dir is validated to stay inside an allowed root before use."""
    if override_company_dir:
        allowed_roots = (os.path.abspath("trailerd"), os.path.abspath("online-platform"))
        resolved = os.path.abspath(override_company_dir)
        if not any(resolved == root or resolved.startswith(root + os.sep) for root in allowed_roots):
            raise HTTPException(status_code=400, detail="Invalid storage location")
    try:
        rerun_record = None
        if rerun_id and rerun_id > 0:
            rerun_record = get_resume_by_id(rerun_id)
            if not rerun_record:
                raise HTTPException(status_code=404, detail="Record to re-run was not found")

        # Validate JD text length
        if not jd_text or len(jd_text) > 50000:
            raise HTTPException(status_code=400, detail="Invalid job description")

        if len(jd_text) < 50:
            append_skipped_to_csv(jd_text, "JD too short (min 50 chars)")
            raise HTTPException(status_code=400, detail="Job description too short (minimum 50 characters)")

        years = _extract_experience_years(jd_text)
        if years > 10:
            reason = f"Requires {years}+ years experience (max allowed: 10)"
            append_skipped_to_csv(jd_text, reason)
            raise HTTPException(status_code=400, detail=reason)

        gc_reason = _check_visa_eligibility(jd_text)
        if gc_reason:
            append_skipped_to_csv(jd_text, gc_reason)
            raise HTTPException(status_code=400, detail=gc_reason)

        lang_reason = _check_foreign_language(jd_text)
        if lang_reason:
            append_skipped_to_csv(jd_text, lang_reason)
            raise HTTPException(status_code=400, detail=lang_reason)

        lead_reason = _check_lead_role(jd_text)
        if lead_reason:
            append_skipped_to_csv(jd_text, lead_reason)
            raise HTTPException(status_code=400, detail=lead_reason)

        if resume:
            file_bytes = await resume.read()
            if len(file_bytes) > 5 * 1024 * 1024:
                raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")
            if not resume.filename or not resume.filename.endswith('.docx'):
                raise HTTPException(status_code=400, detail="Only .docx files are supported")
            if len(file_bytes) == 0:
                raise HTTPException(status_code=400, detail="File is empty")

            safe_name = re.sub(r'[^\w\-.]', '_', os.path.basename(resume.filename))
            safe_name = re.sub(r"\.{2,}", ".", safe_name)  # Prevent directory traversal
            original_filename = os.path.join("original", safe_name)

            # Verify path is within original directory
            abs_path = os.path.abspath(original_filename)
            abs_base = os.path.abspath("original")
            if not abs_path.startswith(abs_base):
                raise HTTPException(status_code=400, detail="Invalid filename")

            with open(original_filename, "wb") as f:
                f.write(file_bytes)
        elif selected_resume:
            safe = os.path.basename(selected_resume)
            resume_path = os.path.join("original", safe)

            # Verify path is within original directory
            abs_path = os.path.abspath(resume_path)
            abs_base = os.path.abspath("original")
            if not abs_path.startswith(abs_base):
                raise HTTPException(status_code=400, detail="Resume not found")

            if not os.path.exists(resume_path):
                raise HTTPException(status_code=400, detail="Resume not found")

            # Check file size before loading
            file_size = os.path.getsize(resume_path)
            if file_size > 5 * 1024 * 1024:
                raise HTTPException(status_code=400, detail="Resume file too large")

            with open(resume_path, "rb") as f:
                file_bytes = f.read()
            original_filename = resume_path
        else:
            base_path = "original/base_resume.docx"
            if not os.path.exists(base_path):
                raise HTTPException(status_code=400, detail="No resume found. Please upload a resume first.")
            with open(base_path, "rb") as f:
                file_bytes = f.read()
            original_filename = base_path

        resume_text = extract_text_from_docx(file_bytes)

        # ── Duplicate JD check — skipped when explicitly re-running an existing record,
        # or when the caller is Command Center tailoring the same job a second time
        # (its own JD naturally matches the scan record it created the first time) ──
        if not rerun_record and not skip_duplicate_check:
            all_history = get_all_resumes(limit=1000)
            for item in all_history:
                if not item.get('jd_text'):
                    continue
                ratio = difflib.SequenceMatcher(None, jd_text, item['jd_text']).ratio()
                if ratio >= 0.95:
                    company = item.get('company_name', 'Unknown')
                    dup_score = item.get('score', 0)
                    logger.info(f"Duplicate JD detected — {ratio:.0%} match with {company} (score={dup_score})")
                    raise HTTPException(
                        status_code=409,
                        detail={
                            "message": f"This JD matches an existing entry for {company} (score: {dup_score}%).",
                            "duplicate_id": item.get('id'),
                            "company_name": company,
                            "score": dup_score,
                        },
                    )

        # ── AI analysis ──
        try:
            result = analyze_resume(resume_text, jd_text, ai_notes=ai_notes or "")
        except RuntimeError as e:
            logger.warning(f"AI provider unavailable: {e}")
            raise HTTPException(status_code=503, detail="The AI provider is currently unavailable. Please try again in a few moments.")

        score = result.get('score', 0)
        company_name = result.get('company_name', _extract_company_name(jd_text))
        contact_info = result.get('contact_info', {})
        jd_emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', jd_text)
        vendor_email = next((e for e in jd_emails if e.split('@')[1].split('.')[0].lower() not in _GENERIC_EMAIL), '')

        if rerun_record and rerun_record.get('file_path'):
            # Re-run: reuse the SAME company directory/file — never create a new entry
            company_dir = os.path.dirname(rerun_record['file_path'])
            os.makedirs(company_dir, exist_ok=True)
        elif override_company_dir:
            # Command Center tailoring: use the job's own deterministic artifact
            # folder instead of an AI-detected company name, which can mismatch the
            # job's actual stored company (e.g. JD text with no explicit company line).
            company_dir = override_company_dir
            os.makedirs(company_dir, exist_ok=True)
        else:
            company_dir = _make_company_dir(company_name, vendor_email, root=storage_root)

        file_path = f"{company_dir}/{RESUME_FILENAME}"
        jd_path = f"{company_dir}/jd_info.txt"
        diff_path = f"{company_dir}/difference.txt"

        replacements = result.get('replacements', [])
        after_score = result.get('after_score', score)

        # ── Decide: skip tailoring, reuse existing, or create new ──
        if score >= 85:
            # Base resume already strong — use original as-is, no tailoring
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            after_score = score  # No tailoring done; keep actual score
            replacements = []
            logger.info(f"Score {score}% >= 85% for {company_name} — no tailoring needed")
            with open(diff_path, "w", encoding="utf-8") as f:
                f.write(f"Score: {score}% — already above 85%%, no tailoring needed.\n")

        elif replacements:
            tailored_stream = create_tailored_docx(file_bytes, replacements)
            with open(file_path, "wb") as f:
                f.write(tailored_stream.read())
            logger.info(f"Applied {len(replacements)} replacements for {company_name}")

            missing_kw = result.get('missing_keywords', [])
            with open(diff_path, "w", encoding="utf-8") as f:
                f.write(f"Resume Tailoring Differences for {company_name}\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"Score: {score}% → {after_score}% (Δ +{after_score - score})\n")
                f.write(f"Total Replacements: {len(replacements)}\n")
                if missing_kw:
                    f.write(f"Missing Keywords: {', '.join(missing_kw)}\n")
                f.write("=" * 60 + "\n\n")
                for i, rep in enumerate(replacements, 1):
                    kw_added = rep.get('keywords_added', [])
                    f.write(f"--- REPLACEMENT #{i} ---\n")
                    if kw_added:
                        f.write(f"Keywords Added: {', '.join(kw_added)}\n\n")
                    f.write(f"REMOVED:\n{rep.get('original', 'N/A')}\n\n")
                    f.write(f"ADDED:\n{rep.get('new', 'N/A')}\n")
                    f.write("=" * 60 + "\n\n")
        else:
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            after_score = score
            with open(diff_path, "w", encoding="utf-8") as f:
                f.write("No changes were needed. The resume already matched the job description well.\n")

        contact_info = result.get('contact_info', {})
        with open(jd_path, "w", encoding="utf-8") as f:
            f.write("Contact Info:\n")
            f.write(f"Name: {contact_info.get('name', 'N/A')}\n")
            f.write(f"Email: {contact_info.get('email', 'N/A')}\n")
            f.write(f"Phone: {contact_info.get('phone', 'N/A')}\n\n")
            f.write("Job Description:\n")
            f.write(jd_text)

        pdf_path = _convert_docx_to_pdf(file_path)

        scan_data = {
            "score": score,
            "after_score": after_score,
            "missing_keywords": result.get('missing_keywords', []),
            "section_scores": result.get('section_scores', {}),
            "contact_info": contact_info,
            "replacements": replacements,
        }

        if rerun_record:
            update_resume_rerun(rerun_record['id'], company_name, jd_text, after_score, json.dumps(scan_data))
            record_id = rerun_record['id']
            existing = None
            logger.info(f"Re-ran scan for record {record_id} ({company_name}) — updated in place")
        else:
            existing = find_existing_company(company_name)
            record_id = save_resume_record(company_name, jd_text, after_score, file_path, json.dumps(scan_data))
        append_to_csv(company_name, jd_text, score, after_score, original_filename, file_path)

        return {
            "id": record_id,
            "company_name": company_name,
            "file_path": file_path,
            "pdf_path": pdf_path,
            "tailored": len(replacements) > 0,
            "duplicate": existing is not None,
            "previous_score": existing['score'] if existing else None,
            "rerun": rerun_record is not None,
            **scan_data,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Scan resume error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to process resume. Please try again later.")


@app.post("/api/scan")
@limiter.limit("10/minute")
async def scan_resume(request: Request, jd_text: str = Form(...), resume: Optional[UploadFile] = File(None), selected_resume: Optional[str] = Form(None), ai_notes: Optional[str] = Form(None), rerun_id: Optional[int] = Form(None)):
    """Public HTTP endpoint — only forwards user-facing fields. storage_root,
    override_company_dir, and skip_duplicate_check are internal-only and can never be
    set by a request to this endpoint; see _scan_resume_core."""
    return await _scan_resume_core(request, jd_text, resume, selected_resume, ai_notes, rerun_id)


def _process_single_jd(jd_text: str, file_bytes: bytes, original_filename: str, ai_notes: str = "") -> dict:
    resume_text = extract_text_from_docx(file_bytes)

    # Duplicate JD check
    all_history = get_all_resumes(limit=1000)
    for item in all_history:
        if not item.get('jd_text'):
            continue
        ratio = difflib.SequenceMatcher(None, jd_text, item['jd_text']).ratio()
        if ratio >= 0.95:
            company = item.get('company_name', 'Unknown')
            return {
                "skipped": True,
                "reason": f"Duplicate JD — already scanned for {company} (score: {item.get('score', 0)}%)",
            }

    result = None
    for attempt in range(3):
        try:
            result = analyze_resume(resume_text, jd_text, ai_notes=ai_notes)
            break
        except RuntimeError as e:
            if attempt < 2 and ("503" in str(e) or "UNAVAILABLE" in str(e) or "429" in str(e) or "All models" in str(e)):
                wait = (attempt + 1) * 10
                logger.warning(f"AI unavailable (attempt {attempt + 1}/3), retrying in {wait}s...")
                time.sleep(wait)
            else:
                raise
    if result is None:
        raise RuntimeError("AI unavailable after 3 attempts")

    score = result.get('score', 0)
    ai_company = result.get('company_name', '')
    local_company = _extract_company_name(jd_text)

    contact_info = result.get('contact_info', {})
    jd_emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', jd_text)
    vendor_email = next((e for e in jd_emails if e.split('@')[1].split('.')[0].lower() not in _GENERIC_EMAIL), '')

    if ai_company and ai_company not in ('Unknown_Company', 'Unknown', ''):
        company_name = ai_company
    elif local_company != 'Unknown_Company':
        company_name = local_company
    elif vendor_email:
        company_name = vendor_email
    else:
        company_name = ai_company or local_company
    company_dir = _make_company_dir(company_name, vendor_email)
    file_path = f"{company_dir}/{RESUME_FILENAME}"
    jd_path = f"{company_dir}/jd_info.txt"
    diff_path = f"{company_dir}/difference.txt"

    replacements = result.get('replacements', [])
    after_score = result.get('after_score', score)

    if score >= 85:
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        after_score = score  # No tailoring done; keep actual score
        replacements = []
        with open(diff_path, "w", encoding="utf-8") as f:
            f.write(f"Score: {score}% — already above 85%, no tailoring needed.\n")

    elif replacements:
        tailored_stream = create_tailored_docx(file_bytes, replacements)
        with open(file_path, "wb") as f:
            f.write(tailored_stream.read())
        missing_kw = result.get('missing_keywords', [])
        with open(diff_path, "w", encoding="utf-8") as f:
            f.write(f"Resume Tailoring Differences for {company_name}\n")
            f.write("=" * 60 + "\n\n")
            f.write(f"Score: {score}% → {after_score}% (Δ +{after_score - score})\n")
            f.write(f"Total Replacements: {len(replacements)}\n")
            if missing_kw:
                f.write(f"Missing Keywords: {', '.join(missing_kw)}\n")
            f.write("=" * 60 + "\n\n")
            for i, rep in enumerate(replacements, 1):
                kw_added = rep.get('keywords_added', [])
                f.write(f"--- REPLACEMENT #{i} ---\n")
                if kw_added:
                    f.write(f"Keywords Added: {', '.join(kw_added)}\n\n")
                f.write(f"REMOVED:\n{rep.get('original', 'N/A')}\n\n")
                f.write(f"ADDED:\n{rep.get('new', 'N/A')}\n")
                f.write("=" * 60 + "\n\n")
    else:
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        after_score = score
        with open(diff_path, "w", encoding="utf-8") as f:
            f.write("No changes were needed. The resume already matched the job description well.\n")

    contact_info = result.get('contact_info', {})
    with open(jd_path, "w", encoding="utf-8") as f:
        f.write("Contact Info:\n")
        f.write(f"Name: {contact_info.get('name', 'N/A')}\n")
        f.write(f"Email: {contact_info.get('email', 'N/A')}\n")
        f.write(f"Phone: {contact_info.get('phone', 'N/A')}\n\n")
        f.write("Job Description:\n")
        f.write(jd_text)

    pdf_path = _convert_docx_to_pdf(file_path)

    existing = find_existing_company(company_name)
    scan_data = {
        "score": score, "after_score": after_score,
        "missing_keywords": result.get('missing_keywords', []),
        "section_scores": result.get('section_scores', {}),
        "contact_info": contact_info, "replacements": replacements,
    }
    record_id = save_resume_record(company_name, jd_text, after_score, file_path, json.dumps(scan_data))
    append_to_csv(company_name, jd_text, score, after_score, original_filename, file_path)
    return {
        "id": record_id, "company_name": company_name, "file_path": file_path,
        "pdf_path": pdf_path,
        "tailored": len(replacements) > 0,
        "duplicate": existing is not None,
        "previous_score": existing['score'] if existing else None, **scan_data,
    }

@app.post("/api/batch-scan")
@limiter.limit("3/minute")
async def batch_scan(request: Request, body: BatchScanRequest):
    try:
        if len(body.jd_texts) > 10:
            raise HTTPException(status_code=400, detail="Maximum 10 JDs per batch")
        if len(body.jd_texts) == 0:
            raise HTTPException(status_code=400, detail="No JDs provided")

        safe = os.path.basename(body.selected_resume)
        resume_path = os.path.join("original", safe)
        abs_path = os.path.abspath(resume_path)
        abs_base = os.path.abspath("original")
        if not abs_path.startswith(abs_base):
            raise HTTPException(status_code=400, detail="Resume not found")
        if not os.path.exists(resume_path):
            raise HTTPException(status_code=400, detail="Resume not found")
        if os.path.getsize(resume_path) > 5 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Resume file too large")

        with open(resume_path, "rb") as f:
            file_bytes = f.read()

        results = []

        for idx, jd_text in enumerate(body.jd_texts):
            if not jd_text or len(jd_text) < 50:
                reason = "JD too short (min 50 chars)"
                results.append({"index": idx, "skipped": True, "reason": reason})
                append_skipped_to_csv(jd_text or "", reason)
                continue
            if len(jd_text) > 50000:
                reason = "JD too long (max 50000 chars)"
                results.append({"index": idx, "skipped": True, "reason": reason})
                append_skipped_to_csv(jd_text[:50000], reason)
                continue

            years = _extract_experience_years(jd_text)
            if years > 10:
                reason = f"Requires {years}+ years experience (max allowed: 10)"
                results.append({"index": idx, "skipped": True, "reason": reason})
                append_skipped_to_csv(jd_text, reason)
                continue

            gc_reason = _check_visa_eligibility(jd_text)
            if gc_reason:
                results.append({"index": idx, "skipped": True, "reason": gc_reason})
                append_skipped_to_csv(jd_text, gc_reason)
                continue

            lang_reason = _check_foreign_language(jd_text)
            if lang_reason:
                results.append({"index": idx, "skipped": True, "reason": lang_reason})
                append_skipped_to_csv(jd_text, lang_reason)
                continue

            lead_reason = _check_lead_role(jd_text)
            if lead_reason:
                results.append({"index": idx, "skipped": True, "reason": lead_reason})
                append_skipped_to_csv(jd_text, lead_reason)
                continue

            try:
                result = _process_single_jd(jd_text, file_bytes, resume_path, ai_notes=body.ai_notes or "")
                result["index"] = idx
                result["skipped"] = False
                results.append(result)
                logger.info(f"Batch JD #{idx + 1} done: {result.get('company_name', 'Unknown')}")
            except Exception as e:
                logger.error(f"Batch scan error for JD #{idx}: {e}", exc_info=True)
                reason = f"Processing failed: {str(e)}"
                results.append({"index": idx, "skipped": True, "reason": reason})
                append_skipped_to_csv(jd_text, reason)

            if idx < len(body.jd_texts) - 1:
                time.sleep(2)
        logger.info(f"Batch scan completed: {len(body.jd_texts)} JDs, {sum(1 for r in results if not r.get('skipped'))} processed, {sum(1 for r in results if r.get('skipped'))} skipped")
        return {"results": results, "total": len(body.jd_texts), "processed": sum(1 for r in results if not r.get('skipped')), "skipped": sum(1 for r in results if r.get('skipped'))}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Batch scan error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Batch processing failed. Please try again later.")

@app.get("/api/history")
@limiter.limit("60/minute")
async def get_history(request: Request, limit: int = 50, offset: int = 0):
    try:
        # Validate input parameters
        limit = min(max(limit, 1), 200)  # Clamp between 1 and 200
        offset = max(offset, 0)
        records = get_all_resumes(limit, offset)
        for r in records:
            fp = r.get('file_path')
            pdf_path = f"{os.path.splitext(fp)[0]}.pdf" if fp else None
            r['pdf_path'] = pdf_path if pdf_path and os.path.exists(pdf_path) else None
        return records
    except Exception as e:
        logger.error(f"Get history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve history")
@app.get("/api/addresses")
@limiter.limit("60/minute")
async def get_addresses(request: Request):
    from database import get_all_addresses
    from services.ollama_service import _extract_recruiter_name
    import re
    try:
        records = get_all_addresses()
        results = []
        for rec in records:
            jd = rec.get("jd_text") or ""
            
            emails_found = set()
            for m in re.finditer(r'[\w.+-]+@[\w-]+\.[\w.-]+', jd):
                emails_found.add(m.group().lower())
                
            emails_list = list(emails_found)
            email = emails_list[0] if emails_list else ""
            
            # Extract phone number with extension if available
            phones = re.findall(r'(?:\+?\d{1,2}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}(?:\s*(?:ext\.?|x|ext\s*:)\s*:?\s*\d+)?', jd, re.IGNORECASE)
            phone = phones[0].strip() if phones else ""
            
            # Fallback to extracting name
            from services.ollama_service import _extract_recruiter_name
            name = _extract_recruiter_name(jd)
            
            results.append({
                "id": rec["id"],
                "company_name": rec["company_name"],
                "user_address": rec["user_address"],
                "phone": phone,
                "name": name,
                "email": email,
                "all_emails": emails_list
            })
        return results
    except Exception as e:
        logger.error(f"Get addresses error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve addresses")

@app.get("/api/search")
@limiter.limit("30/minute")
async def search_history(request: Request, q: str = "", limit: int = 50):
    if not q or len(q.strip()) < 2:
        raise HTTPException(status_code=400, detail="Query must be at least 2 characters")
    try:
        limit = min(max(limit, 1), 100)
        records = search_records(q.strip(), limit)
        results = []
        for rec in records:
            jd = rec.get("jd_text") or ""
            sr = rec.get("scan_result") or {}
            contact = sr.get("contact_info") or {}

            emails_found = set()
            for m in re.finditer(r'[\w.+-]+@[\w-]+\.[\w.-]+', jd):
                emails_found.add(m.group().lower())

            jd_lower = jd.lower()
            location = "Not specified"
            loc_m = re.search(r'(?:location|city|office|based in|work location)[:\s]*([^\n]{3,60})', jd, re.IGNORECASE)
            if loc_m:
                location = loc_m.group(1).strip().rstrip('.,:;')

            local_required = False
            local_keywords = ["local candidates", "locals only", "local only", "must be local",
                              "onsite only", "on-site only", "no relocation", "local to",
                              "must reside", "must live in", "within commuting"]
            for kw in local_keywords:
                if kw in jd_lower:
                    local_required = True
                    break

            position = "Not specified"
            role_patterns = [
                r'^(?:job\s+title|position\s+title|role\s+title)\s*[:\-–]\s*(.{3,80})$',
                r'^(?:title|position|role)\s*:\s*(.{3,80})$',
                r'^(?:job|opening|vacancy)\s*:\s*(.{3,80})$',
                r'(?:hiring\s+(?:for|a)|looking\s+for\s+(?:a|an))\s+([A-Z][A-Za-z /&\-,]+)',
            ]
            for pat in role_patterns:
                role_m = re.search(pat, jd, re.IGNORECASE | re.MULTILINE)
                if role_m:
                    candidate = role_m.group(1).strip().rstrip('.,:;')
                    if len(candidate) >= 3 and len(candidate) <= 80:
                        position = candidate
                        break
            if position == "Not specified":
                first_line = jd.strip().split('\n')[0].strip() if jd.strip() else ""
                if 5 <= len(first_line) <= 80 and not any(w in first_line.lower() for w in ['hi ', 'hello', 'dear ', 'please', 'hope ', 'i am']):
                    position = first_line.rstrip('.,:;')

            recruiter_name = None
            name_m = re.search(r'(?:recruiter|hiring manager|contact|regards|thanks|best|sincerely)[,:\s]*\n?\s*([A-Z][a-z]+ [A-Z][a-z]+)', jd)
            if name_m:
                recruiter_name = name_m.group(1).strip()

            results.append({
                "id": rec["id"],
                "company_name": rec.get("company_name", "Unknown"),
                "position": position,
                "emails": list(emails_found),
                "recruiter_name": recruiter_name,
                "location": location,
                "local_required": local_required,
                "user_address": rec.get("user_address") or "",
                "user_notes": rec.get("user_notes") or "",
                "status": rec.get("status", "Scanned"),
                "score": rec.get("score", 0),
                "match_percentage": rec.get("match_percentage"),
                "source": rec.get("source", "dashboard"),
                "created_at": rec.get("created_at"),
                "jd_preview": jd[:200] if jd else "",
            })
        return {"results": results, "count": len(results)}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Search error: {e}")
        raise HTTPException(status_code=500, detail="Search failed")


class AddressUpdate(BaseModel):
    address: str

@app.patch("/api/history/{record_id}/address")
@limiter.limit("30/minute")
async def patch_user_address(request: Request, record_id: int, body: AddressUpdate):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")
        if len(body.address) > 500:
            raise HTTPException(status_code=400, detail="Address too long")
        update_user_address(record_id, body.address.strip())
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update address error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update address")

class NotesUpdate(BaseModel):
    notes: str

@app.patch("/api/history/{record_id}/notes")
@limiter.limit("30/minute")
async def patch_user_notes(request: Request, record_id: int, body: NotesUpdate):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")
        if len(body.notes) > 2000:
            raise HTTPException(status_code=400, detail="Notes too long")
        update_user_notes(record_id, body.notes.strip())
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update notes error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update notes")

@app.delete("/api/history/{record_id}")
@limiter.limit("20/minute")
async def delete_history_item(request: Request, record_id: int):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")
        delete_resume_record(record_id)
        logger.info(f"History record deleted: {record_id}")
        return {"status": "ok", "message": "Record deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete record")

@app.patch("/api/history/{record_id}/status")
@limiter.limit("20/minute")
async def patch_history_status(request: Request, record_id: int, status_update: StatusUpdate):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")
        # Validate status value
        valid_statuses = ["Scanned", "Applied", "Phone Screen", "Interview", "Offer", "Rejected"]
        if status_update.status not in valid_statuses:
            raise HTTPException(status_code=400, detail=f"Invalid status. Must be one of: {', '.join(valid_statuses)}")
        update_resume_status(record_id, status_update.status)
        logger.info(f"History record status updated: {record_id} -> {status_update.status}")
        return {"status": "ok", "message": "Status updated"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Update status error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update status")

@app.post("/api/history/{record_id}/cover-letter")
@limiter.limit("5/minute")
async def api_generate_cover_letter(request: Request, record_id: int):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        file_path = _resolve_resume_path(record)
        if not file_path:
            raise HTTPException(status_code=400, detail="Resume file not found")

        # Verify file path is valid
        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            raise HTTPException(status_code=400, detail="Resume file not accessible")

        with open(abs_path, "rb") as f:
            resume_text = extract_text_from_docx(f.read())

        # Command Center records store only the job title in jd_text — the full
        # description lives in scan_result['description']. Prefer that when present
        # so the cover letter is written against the real JD, not just the title.
        scan_result = record.get('scan_result')
        command_center_description = (scan_result.get('description') or '').strip() if isinstance(scan_result, dict) else ''
        jd_text = command_center_description if len(command_center_description) >= 50 else record['jd_text']

        cover_letter = generate_cover_letter(resume_text, jd_text, record['company_name'])

        company_dir = os.path.dirname(abs_path)
        cl_filename = get_output_filename("cover_letter", record_id)
        cl_path = os.path.join(company_dir, cl_filename).replace("\\", "/")

        cl_doc = docx.Document()
        cl_doc.add_paragraph(cover_letter)
        cl_doc.save(cl_path)

        from database import mark_cover_letter_generated
        mark_cover_letter_generated(record_id)

        logger.info(f"Cover letter generated for record: {record_id}")
        return {"cover_letter": cover_letter, "cl_path": cl_path}
    except RuntimeError as e:
        logger.warning(f"AI provider error: {e}")
        raise HTTPException(status_code=503, detail="The AI provider is unavailable. Please try again in a few moments.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Cover letter generation error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate cover letter")

class AddPointsRequest(BaseModel):
    points: str
    target_hint: Optional[str] = None

@app.post("/api/history/{record_id}/add-points")
@limiter.limit("10/minute")
async def api_add_points(request: Request, record_id: int, payload: AddPointsRequest):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        points_text = (payload.points or "").strip()
        if not points_text:
            raise HTTPException(status_code=400, detail="Please enter at least one point to add")
        if len(points_text) > 5000:
            raise HTTPException(status_code=400, detail="Points text too long (maximum 5000 characters)")

        target_hint = (payload.target_hint or "").strip()
        if len(target_hint) > 200:
            raise HTTPException(status_code=400, detail="Target project/company name too long")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        file_path = _resolve_resume_path(record)
        if not file_path:
            raise HTTPException(status_code=400, detail="Tailored resume file not found for this record")

        with open(file_path, "rb") as f:
            current_bytes = f.read()
        resume_text = extract_text_from_docx(current_bytes)

        try:
            ai_result = generate_additional_points(resume_text, record.get('jd_text', ''), points_text, target_hint)
        except RuntimeError as e:
            logger.warning(f"AI provider unavailable: {e}")
            raise HTTPException(status_code=503, detail="The AI provider is currently unavailable. Please try again in a few moments.")

        insertions = ai_result.get('insertions', [])
        if not insertions:
            raise HTTPException(status_code=422, detail="Could not determine where to add these points. Try rephrasing or specifying a project/company name.")

        updated_stream = insert_bullets_after(current_bytes, insertions)
        # Update the SAME tailored resume file in place — never create a new file or record
        with open(file_path, "wb") as f:
            f.write(updated_stream.read())

        existing_scan = record.get('scan_result') or {}
        if not isinstance(existing_scan, dict):
            existing_scan = {}
        existing_replacements = existing_scan.get('replacements', [])

        added_entries = [
            {
                "original": f"[Added under {ins.get('section', 'resume')}]",
                "new": ins.get('new_bullet', ''),
                "keywords_added": ins.get('keywords_added', []),
            }
            for ins in insertions if ins.get('new_bullet')
        ]

        new_after_score = ai_result.get('after_score', existing_scan.get('after_score', record.get('score')))

        merged_scan = {
            **existing_scan,
            "after_score": new_after_score,
            "replacements": existing_replacements + added_entries,
        }

        update_resume_after_edit(record_id, new_after_score, json.dumps(merged_scan))

        diff_path = os.path.join(os.path.dirname(file_path), "difference.txt").replace("\\", "/")
        try:
            with open(diff_path, "a", encoding="utf-8") as f:
                f.write(f"\n--- ADDITIONAL POINTS ({datetime.now().isoformat()}) ---\n")
                for ins in insertions:
                    f.write(f"Section: {ins.get('section', 'N/A')}\n")
                    f.write(f"ADDED:\n{ins.get('new_bullet', 'N/A')}\n\n")
        except OSError:
            pass

        logger.info(f"Added {len(insertions)} point(s) to record {record_id}")
        return {
            "id": record_id,
            "inserted": len(insertions),
            "insertions": insertions,
            "scan_result": merged_scan,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Add points error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to add points to the tailored resume")

class MailDraftSave(BaseModel):
    subject: str
    body: str

@app.post("/api/history/{record_id}/mail-draft")
@limiter.limit("5/minute")
async def api_generate_mail_draft(request: Request, record_id: int):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        file_path = _resolve_resume_path(record)
        if not file_path:
            raise HTTPException(status_code=400, detail="No file path for this record")

        company_dir = os.path.dirname(file_path)

        # Read ALL files in the company folder for full context
        resume_text = ""
        jd_text = ""
        cover_letter_text = ""
        jd_info_text = ""

        # Read the tailored resume using the exact path from the database
        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                resume_text = extract_text_from_docx(f.read())

        # Read other files from the company folder
        if os.path.isdir(company_dir):
            for filename in os.listdir(company_dir):
                filepath = os.path.join(company_dir, filename)
                if filename == "jd_info.txt":
                    with open(filepath, "r", encoding="utf-8") as f:
                        jd_info_text = f.read()
                elif filename.endswith(".docx") and filepath != file_path:
                    # Read cover letter (any docx that isn't the resume)
                    if "cover" in filename.lower() or filename.startswith("cover_letter_"):
                        with open(filepath, "rb") as f:
                            cover_letter_text = extract_text_from_docx(f.read())

        if not resume_text:
            raise HTTPException(status_code=400, detail="Could not read resume text")

        # Use full JD from database (most complete), fall back to jd_info.txt
        jd_text = record.get('jd_text', '') or jd_info_text
        if not jd_text:
            raise HTTPException(status_code=400, detail="No job description found for this record")

        # Read personal profile if it exists (work auth, location, etc.)
        profile_text = ""
        profile_path = os.path.join(DATA_DIR, "profile.txt")
        if os.path.exists(profile_path):
            with open(profile_path, "r", encoding="utf-8") as f:
                profile_text = f.read().strip()

        # Extract any email addresses found in the JD
        to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', jd_text)))

        mail = generate_mail_draft(resume_text, jd_text, cover_letter_text, record['company_name'], profile_text)
        return {
            "to_emails": to_emails,
            "subject": mail.get('subject', ''),
            "body": mail.get('body', ''),
        }
    except RuntimeError as e:
        logger.warning(f"AI provider error: {e}")
        raise HTTPException(status_code=503, detail="The AI provider is unavailable. Please try again in a few moments.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Mail draft generation error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate mail draft")

@app.post("/api/history/{record_id}/mail-draft/save")
@limiter.limit("10/minute")
async def save_mail_draft(request: Request, record_id: int, draft: MailDraftSave):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        resolved_path = _resolve_resume_path(record)
        raw_file_path = resolved_path or record.get('file_path')
        if not raw_file_path:
            raise HTTPException(status_code=400, detail="No file path for this record")

        company_dir = os.path.dirname(raw_file_path)
        os.makedirs(company_dir, exist_ok=True)
        draft_filename = get_output_filename("mail_draft", record_id)
        draft_path = os.path.join(company_dir, draft_filename).replace("\\", "/")

        content = f"Subject: {draft.subject}\n\n{draft.body}"
        with open(draft_path, "w", encoding="utf-8") as f:
            f.write(content)

        logger.info(f"Mail draft saved for record: {record_id}")
        return {"draft_path": draft_path}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Save mail draft error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save mail draft")

@app.get("/api/history/{record_id}/content")
@limiter.limit("60/minute")
async def get_record_content(request: Request, record_id: int):
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        file_path = _resolve_resume_path(record) or record.get('file_path', '')
        if not file_path:
            raise HTTPException(status_code=400, detail="No file path for this record")

        company_dir = os.path.dirname(file_path)
        payload = {"company_name": record['company_name'], "cover_letter": None, "cl_path": None, "mail_draft": None, "draft_path": None, "follow_up_draft": None}

        # Find cover letter files
        cl_files = [f for f in os.listdir(company_dir) if f.startswith("cover_letter_") and f.endswith(".docx")] if os.path.isdir(company_dir) else []
        if cl_files:
            try:
                cl_full_path = os.path.join(company_dir, cl_files[0])
                with open(cl_full_path, "rb") as f:
                    payload["cover_letter"] = extract_text_from_docx(f.read())
                payload["cl_path"] = cl_full_path.replace("\\", "/")
            except Exception:
                pass

        # Find mail draft files
        draft_files = [f for f in os.listdir(company_dir) if f.startswith("mail_draft_") and f.endswith(".txt")] if os.path.isdir(company_dir) else []
        if draft_files:
            try:
                draft_full_path = os.path.join(company_dir, draft_files[0])
                with open(draft_full_path, "r", encoding="utf-8") as f:
                    content = f.read()
                if content.startswith("Subject:"):
                    sep = content.find('\n\n')
                    subject = content[len("Subject:"):sep].strip() if sep != -1 else content[len("Subject:"):].strip()
                    body = content[sep + 2:] if sep != -1 else ""
                else:
                    subject, body = "", content
                to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', record.get('jd_text', ''))))
                payload["mail_draft"] = {"to_emails": to_emails, "subject": subject, "body": body}
                payload["draft_path"] = draft_full_path.replace("\\", "/")
            except Exception:
                pass

        # Load saved follow-up draft from DB
        try:
            saved_fu = get_follow_up_draft(record_id)
            if saved_fu:
                payload["follow_up_draft"] = saved_fu
            elif os.path.isdir(company_dir):
                fu_file = os.path.join(company_dir, "follow_up_draft.txt")
                if os.path.exists(fu_file):
                    with open(fu_file, "r", encoding="utf-8") as f:
                        content = f.read()
                    sep = content.find('\n\n')
                    fu_subject = content[len("Subject:"):sep].strip() if content.startswith("Subject:") and sep != -1 else ""
                    fu_body = content[sep + 2:] if sep != -1 else content
                    payload["follow_up_draft"] = {"subject": fu_subject, "body": fu_body}
        except Exception:
            pass

        return payload
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get record content error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve record content")

@app.get("/api/history/csv")
@limiter.limit("10/minute")
async def download_history_csv(request: Request):
    try:
        csv_file = os.path.join(DATA_DIR, "history.csv")
        if not os.path.exists(csv_file):
            raise HTTPException(status_code=404, detail="No history CSV found yet.")
        logger.info("History CSV downloaded")
        return FileResponse(csv_file, filename="resume_history.csv", media_type="text/csv")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download CSV error: {e}")
        raise HTTPException(status_code=500, detail="Failed to download CSV")

@app.get("/api/download/{file_path:path}")
def download_resume(file_path: str):
    try:
        # Security: allow subdirectories (e.g. CompanyName/resume.docx) but
        # prevent path traversal by resolving and checking against base_path.
        # Files live under one of two roots: "trailerd" (Dashboard tailoring) or
        # "online-platform" (Command Center tailoring). Existing frontend callers only
        # ever strip a leading "trailerd/" prefix, so a path still carrying an
        # "online-platform/" prefix is resolved against that root instead; anything with
        # neither prefix falls back to the legacy "trailerd" default.
        rel = file_path.lstrip('/').replace('\\', '/')
        first_segment = rel.split('/', 1)[0]
        if first_segment in ("trailerd", "online-platform"):
            root = first_segment
            rel = rel[len(first_segment) + 1:]
        else:
            root = "trailerd"
        full_path = os.path.abspath(os.path.join(root, rel))
        base_path = os.path.abspath(root)

        # Verify path stays inside the resolved root
        if not full_path.startswith(base_path + os.sep):
            raise HTTPException(status_code=404, detail="File not found")

        if not os.path.exists(full_path) or os.path.isdir(full_path):
            raise HTTPException(status_code=404, detail="File not found")

        filename = os.path.basename(full_path)
        logger.info(f"File downloaded: {rel}")
        return FileResponse(full_path, filename=filename)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Resume download error: {e}")
        raise HTTPException(status_code=500, detail="Failed to download file")

# ─── API Usage Stats ───

@app.get("/api/usage")
async def api_usage():
    """Get API usage statistics and costs."""
    try:
        return get_usage_stats()
    except Exception as e:
        logger.error(f"Usage stats error: {e}")
        return {"today": {"calls": 0, "cost": 0}, "week": {"calls": 0, "cost": 0},
                "month": {"calls": 0, "cost": 0}, "all_time": {"calls": 0, "cost": 0}}

# ─── Personal Profile ───

class ProfileUpdate(BaseModel):
    content: str

@app.get("/api/profile")
async def get_profile():
    """Get the personal profile text."""
    profile_path = os.path.join(DATA_DIR, "profile.txt")
    if os.path.exists(profile_path):
        with open(profile_path, "r", encoding="utf-8") as f:
            return {"content": f.read(), "exists": True}
    return {"content": "", "exists": False}

@app.post("/api/profile")
@limiter.limit("10/minute")
async def save_profile(request: Request, profile: ProfileUpdate):
    """Save personal profile — key facts only, never raw documents."""
    try:
        profile_path = os.path.join(DATA_DIR, "profile.txt")
        with open(profile_path, "w", encoding="utf-8") as f:
            f.write(profile.content)
        logger.info("Personal profile updated")
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"Save profile error: {e}")
        raise HTTPException(status_code=500, detail="Failed to save profile")

@app.post("/api/profile/upload")
@limiter.limit("5/minute")
async def upload_profile_doc(request: Request, file: UploadFile = File(...)):
    """Upload a personal document (PDF, DOCX, image) to extract job-relevant facts.
    Raw file is never stored — only extracted facts are saved to profile."""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        ext = file.filename.lower().rsplit('.', 1)[-1] if '.' in file.filename else ''
        allowed = {'pdf', 'docx', 'png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp'}
        if ext not in allowed:
            raise HTTPException(status_code=400, detail=f"Unsupported file type. Use: {', '.join(allowed)}")

        file_bytes = await file.read()
        if len(file_bytes) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File too large. Maximum 10MB.")
        if len(file_bytes) == 0:
            raise HTTPException(status_code=400, detail="File is empty")

        result = process_uploaded_doc(file_bytes, file.filename)
        logger.info(f"Profile doc processed: {file.filename}")
        return result

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Profile upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to process document")

# ─── Gmail Integration ───

@app.get("/api/gmail/status")
async def gmail_status():
    """Check if Gmail is connected."""
    try:
        return gmail_service.is_connected()
    except Exception as e:
        logger.error(f"Gmail status check error: {e}")
        return {"connected": False}

@app.get("/api/gmail/auth")
async def gmail_auth():
    """Start Gmail OAuth flow — redirects user to Google consent screen."""
    try:
        auth_url = gmail_service.get_auth_url()
        return RedirectResponse(url=auth_url)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail auth error: {e}")
        raise HTTPException(status_code=500, detail="Failed to start Gmail authentication")

@app.get("/api/gmail/callback")
async def gmail_callback(code: str):
    """Handle OAuth callback from Google."""
    try:
        result = gmail_service.handle_callback(code)
        logger.info(f"Gmail connected: {result.get('email', 'unknown')}")
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
        return RedirectResponse(url=f"{frontend_url}?gmail=connected")
    except Exception as e:
        logger.error(f"Gmail callback error: {e}")
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
        return RedirectResponse(url=f"{frontend_url}?gmail=error")

@app.post("/api/gmail/disconnect")
async def gmail_disconnect():
    """Disconnect Gmail account."""
    try:
        result = gmail_service.disconnect()
        logger.info("Gmail disconnected")
        return result
    except Exception as e:
        logger.error(f"Gmail disconnect error: {e}")
        raise HTTPException(status_code=500, detail="Failed to disconnect Gmail")

class GmailDraftRequest(BaseModel):
    to_emails: list = []
    subject: str
    body: str
    record_id: Optional[int] = None
    attach_resume: bool = False
    attach_cover_letter: bool = False
    attach_dl: bool = False
    attach_gc: bool = False

@app.post("/api/gmail/save-draft")
@limiter.limit("10/minute")
async def gmail_save_draft(request: Request, draft: GmailDraftRequest):
    """Save email as a draft in Gmail with selected attachments."""
    try:
        has_selections = draft.attach_resume or draft.attach_cover_letter or draft.attach_dl or draft.attach_gc
        attachment_list = []

        if has_selections:
            if draft.record_id and (draft.attach_resume or draft.attach_cover_letter):
                record = get_resume_by_id(draft.record_id)
                resolved_path = _resolve_resume_path(record) if record else None
                if resolved_path:
                    company_dir = os.path.dirname(resolved_path)
                    if draft.attach_resume:
                        attachment_list.append({
                            "path": resolved_path,
                            "display_name": RESUME_FILENAME
                        })
                    if draft.attach_cover_letter:
                        cl_files = [f for f in os.listdir(company_dir)
                                     if f.endswith(".docx") and ("cover" in f.lower() or f.startswith("cover_letter_"))
                                     and f != os.path.basename(resolved_path)]
                        if cl_files:
                            attachment_list.append({
                                "path": os.path.join(company_dir, cl_files[0]),
                                "display_name": "Teja_Mahesh_Neerukonda_Cover_Letter.docx"
                            })

            doc_map = {"dl": draft.attach_dl, "gc": draft.attach_gc}
            display_names = {"dl": "Drivers_License", "gc": "Green_Card"}
            for doc_type, should_attach in doc_map.items():
                if should_attach:
                    for f in os.listdir(DOCUMENTS_DIR):
                        if f.startswith(f"{doc_type}."):
                            ext = os.path.splitext(f)[1]
                            attachment_list.append({
                                "path": os.path.join(DOCUMENTS_DIR, f),
                                "display_name": f"{display_names[doc_type]}{ext}"
                            })
                            break

        if has_selections:
            result = gmail_service.save_draft(
                to_emails=draft.to_emails,
                subject=draft.subject,
                body=draft.body,
                attachments=attachment_list,
            )
        else:
            attachment_path = None
            if draft.record_id:
                record = get_resume_by_id(draft.record_id)
                if record and record.get('file_path') and os.path.exists(record['file_path']):
                    attachment_path = record['file_path']
            result = gmail_service.save_draft(
                to_emails=draft.to_emails,
                subject=draft.subject,
                body=draft.body,
                attachment_path=attachment_path,
            )

        logger.info(f"Gmail draft saved: {result.get('draft_id', 'unknown')}")
        return result
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail save draft error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to save draft to Gmail")

# ─── Gmail Inbox (for follow-up) ───

@app.get("/api/gmail/inbox/filters")
@limiter.limit("60/minute")
async def gmail_inbox_filters(request: Request):
    return {"filters": gmail_service.list_inbox_filters()}

@app.get("/api/gmail/inbox")
@limiter.limit("60/minute")
async def gmail_inbox(request: Request, q: str = "", category: str = "all", limit: int = 25,
                       page_token: str = None, mode: str = "smart"):
    """Search Gmail messages, then have AI actually read each subject/snippet and
    classify it — one single batched Gemini call for the whole page (not one call per
    email), so accuracy is far better than local keyword rules while staying a fraction
    of a cent per load. Results are cached per message id + content hash (see
    services/inbox_cache.py), so re-visiting a filter after the first read costs nothing
    and feels instant. mode="cheap" skips the AI call entirely and uses Gmail's keyword
    query + local rules only, for zero-API-cost browsing. Falls back to the local-rule
    category for any message the AI call couldn't classify (or if the AI call fails
    entirely — never blocks the inbox from loading). If a specific category was
    requested, the AI's classification (the more accurate read) is used to do the actual
    filtering, not just Gmail's keyword query — and the Gmail query itself now casts a
    broad "job mail" net for any category rather than that category's own narrow
    keywords, so differently-worded messages aren't silently dropped before AI ever sees
    them."""
    try:
        max_results = max(1, min(limit, 50))
        requested_category = (category or "all").lower()
        search_result = gmail_service.search_inbox(q, max_results=max_results, category=requested_category, page_token=page_token)
        messages = search_result["messages"]
        next_page_token = search_result.get("next_page_token")

        if messages and mode != "cheap":
            try:
                from services.ai_service import classify_inbox_messages
                from services import inbox_cache
                cached, uncached = inbox_cache.split_cached(messages)
                ai_categories = dict(cached)
                if uncached:
                    fresh = classify_inbox_messages(uncached)
                    ai_categories.update(fresh)
                    inbox_cache.store(uncached, fresh)
            except Exception as e:
                logger.warning(f"Inbox AI classification failed, keeping local-rule categories: {e}")
                ai_categories = {}

            if ai_categories:
                label_by_key = {f['key']: f['label'] for f in gmail_service.list_inbox_filters()}
                for m in messages:
                    ai_cat = ai_categories.get(m['id'])
                    if ai_cat:
                        m['category'] = ai_cat
                        m['category_label'] = label_by_key.get(ai_cat, ai_cat.title())
                        m['classified_by'] = 'ai'

                # Only re-filter using the AI's (more accurate) read when it actually ran —
                # if it failed, stick with whatever Gmail's own keyword query already
                # returned rather than risk dropping results based on the cruder local
                # rule's category guess.
                if requested_category == "needs_attention":
                    from services.gmail_service import NEEDS_ATTENTION_CATEGORIES
                    messages = [m for m in messages if m['category'] in NEEDS_ATTENTION_CATEGORIES]
                elif requested_category != "all":
                    messages = [m for m in messages if m['category'] == requested_category]
        elif requested_category == "needs_attention":
            # Cheap mode has no AI classification to filter on — fall back to whatever
            # the local-rule category already guessed.
            from services.gmail_service import NEEDS_ATTENTION_CATEGORIES
            messages = [m for m in messages if m['category'] in NEEDS_ATTENTION_CATEGORIES]

        if messages:
            # Cross-reference with the Command Center pipeline — no extra API cost, just
            # a text match — so the Inbox can point back to the tracked application and
            # offer a follow-up action right from the message.
            try:
                from database import get_active_applications
                from services.inbox_matcher import match_message_to_application
                applications = get_active_applications()
                if applications:
                    for m in messages:
                        app = match_message_to_application(m, applications)
                        if app:
                            m['matched_application'] = app
            except Exception as e:
                logger.warning(f"Inbox application matching failed: {e}")

        return {
            "messages": messages,
            "filters": gmail_service.list_inbox_filters(),
            "category": requested_category,
            "next_page_token": next_page_token,
            "mode": mode,
        }
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail inbox error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to search inbox: {e}")

@app.get("/api/gmail/message/{message_id}")
@limiter.limit("20/minute")
async def gmail_message(request: Request, message_id: str):
    """Get full body of a Gmail message."""
    try:
        msg = gmail_service.get_message_body(message_id)
        return msg
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail message read error: {e}")
        raise HTTPException(status_code=500, detail="Failed to read message")

@app.get("/api/gmail/thread/{thread_id}")
@limiter.limit("20/minute")
async def gmail_thread(request: Request, thread_id: str):
    """Full conversation for a message's thread — recruiter back-and-forth (interview
    scheduling, assessment follow-ups, offer negotiation) reads much better as a
    conversation than a single isolated message."""
    try:
        messages = gmail_service.get_thread_messages(thread_id)
        return {"thread_id": thread_id, "messages": messages}
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail thread read error: {e}")
        raise HTTPException(status_code=500, detail="Failed to read thread")

@app.get("/api/gmail/message/{message_id}/summary")
@limiter.limit("20/minute")
async def gmail_message_summary(request: Request, message_id: str, record_id: Optional[int] = None):
    """Full-body AI read of ONE already-opened email — what happened, the required
    action, any deadline/interview date, a recruiter email, reply intent, and a draft
    reply. Reads the WHOLE thread (not just this one message — replies often only make
    sense in light of earlier messages) and, if the message was matched to a tracked
    application (record_id), grounds the reply in that job's title/company too.
    User-triggered only (never run across the inbox list, which stays subject/snippet
    classification for cost), so this is always exactly one cheap call, not one per
    email."""
    try:
        msg = gmail_service.get_message_body(message_id)

        thread_context = ""
        if msg.get('thread_id'):
            try:
                thread_msgs = gmail_service.get_thread_messages(msg['thread_id'])
                earlier = [t for t in thread_msgs if t['id'] != message_id]
                if earlier:
                    thread_context = "\n\n---\n\n".join(
                        f"From: {t['from']}\nDate: {t['date']}\n{t['body'][:1500]}" for t in earlier
                    )
            except Exception as e:
                logger.warning(f"Failed to fetch thread context for summary: {e}")

        application_context = ""
        if record_id:
            try:
                from database import get_job_detail
                job = get_job_detail(record_id)
                if job:
                    application_context = f"{job.get('title', '')} at {job.get('company', '')}"
            except Exception as e:
                logger.warning(f"Failed to load application context for summary: {e}")

        from services.ai_service import summarize_inbox_message
        summary = summarize_inbox_message(
            msg.get('subject', ''), msg.get('from', ''), msg.get('body', ''),
            thread_context=thread_context, application_context=application_context,
        )
        return summary
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail message summary error: {e}")
        raise HTTPException(status_code=500, detail="Failed to summarize message")


class LabelMessageRequest(BaseModel):
    category: str

@app.post("/api/gmail/message/{message_id}/label")
@limiter.limit("30/minute")
async def gmail_label_message(request: Request, message_id: str, payload: LabelMessageRequest):
    """Apply the Job/<Category> Gmail label — requires the gmail.modify scope granted
    on reconnect; raises a clear reconnect prompt (400) if the connected account
    doesn't have it yet, same pattern as the read-permission check on search_inbox."""
    try:
        return gmail_service.apply_job_label(message_id, payload.category)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail label error: {e}")
        raise HTTPException(status_code=500, detail="Failed to apply label")

@app.post("/api/gmail/message/{message_id}/archive")
@limiter.limit("30/minute")
async def gmail_archive_message(request: Request, message_id: str):
    try:
        return gmail_service.archive_message(message_id)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail archive error: {e}")
        raise HTTPException(status_code=500, detail="Failed to archive message")

@app.post("/api/gmail/message/{message_id}/mark-read")
@limiter.limit("30/minute")
async def gmail_mark_read(request: Request, message_id: str):
    try:
        return gmail_service.mark_message_read(message_id)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Gmail mark-read error: {e}")
        raise HTTPException(status_code=500, detail="Failed to mark message read")

# ─── Personal Documents (DL, GC, etc.) ───

DOCUMENTS_DIR = os.path.join(DATA_DIR, "documents")
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

ALLOWED_DOC_TYPES = {"dl", "gc"}

@app.post("/api/documents/upload")
@limiter.limit("10/minute")
async def upload_document(request: Request, doc_type: str = Form(...), file: UploadFile = File(...)):
    """Upload a personal document (DL or GC)."""
    doc_type = doc_type.lower()
    if doc_type not in ALLOWED_DOC_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid document type. Allowed: {', '.join(ALLOWED_DOC_TYPES)}")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"}:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, DOC, PNG, or JPG.")

    save_name = f"{doc_type}{ext}"
    for existing in os.listdir(DOCUMENTS_DIR):
        if existing.startswith(f"{doc_type}."):
            os.remove(os.path.join(DOCUMENTS_DIR, existing))

    save_path = os.path.join(DOCUMENTS_DIR, save_name)
    content = await file.read()
    with open(save_path, "wb") as f:
        f.write(content)

    return {"status": "uploaded", "doc_type": doc_type, "filename": save_name}


@app.get("/api/documents")
async def list_documents(request: Request):
    """List uploaded personal documents."""
    docs = {}
    for f in os.listdir(DOCUMENTS_DIR):
        name = os.path.splitext(f)[0].lower()
        if name in ALLOWED_DOC_TYPES:
            docs[name] = {"filename": f, "path": os.path.join(DOCUMENTS_DIR, f)}
    return {"documents": docs}


@app.delete("/api/documents/{doc_type}")
async def delete_document(request: Request, doc_type: str):
    """Delete a personal document."""
    doc_type = doc_type.lower()
    for f in os.listdir(DOCUMENTS_DIR):
        if f.startswith(f"{doc_type}."):
            os.remove(os.path.join(DOCUMENTS_DIR, f))
            return {"status": "deleted", "doc_type": doc_type}
    raise HTTPException(status_code=404, detail="Document not found")

# ─── Follow-Up Mail ───

class FollowUpRequest(BaseModel):
    received_email: str
    instructions: Optional[str] = None

@app.post("/api/history/{record_id}/follow-up")
@limiter.limit("5/minute")
async def api_generate_follow_up(request: Request, record_id: int, body: FollowUpRequest):
    """Generate a follow-up reply based on a received email."""
    try:
        if record_id <= 0:
            raise HTTPException(status_code=400, detail="Invalid record ID")

        record = get_resume_by_id(record_id)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")

        file_path = _resolve_resume_path(record)
        resume_text = ""
        if file_path and os.path.exists(file_path):
            with open(file_path, "rb") as f:
                resume_text = extract_text_from_docx(f.read())

        jd_text = record.get('jd_text', '')
        if not jd_text:
            raise HTTPException(status_code=400, detail="No job description found for this record")

        # Read existing mail draft if available
        original_mail_body = ""
        company_dir = os.path.dirname(file_path) if file_path else ""
        if company_dir and os.path.isdir(company_dir):
            draft_files = [f for f in os.listdir(company_dir) if f.startswith("mail_draft_") and f.endswith(".txt")]
            if draft_files:
                with open(os.path.join(company_dir, draft_files[0]), "r", encoding="utf-8") as f:
                    original_mail_body = f.read()

        profile_text = ""
        profile_path = os.path.join(DATA_DIR, "profile.txt")
        if os.path.exists(profile_path):
            with open(profile_path, "r", encoding="utf-8") as f:
                profile_text = f.read().strip()

        to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', body.received_email)))

        is_w2 = detect_w2_fulltime(body.received_email)

        conversation_history = ""
        if is_w2 and to_emails:
            try:
                sender_email = to_emails[0]
                convos = gmail_service.get_conversation_with_sender(sender_email, max_results=20)
                if convos:
                    parts = []
                    for msg in convos:
                        parts.append(
                            f"--- Email ({msg['date']}) ---\n"
                            f"From: {msg['from']}\n"
                            f"To: {msg['to']}\n"
                            f"Subject: {msg['subject']}\n\n"
                            f"{msg['body']}"
                        )
                    conversation_history = "\n\n".join(parts)
            except Exception as e:
                logger.warning(f"Could not fetch conversation history: {e}")

        result = generate_follow_up(
            resume_text=resume_text,
            jd_text=jd_text,
            company_name=record['company_name'],
            received_email=body.received_email,
            original_mail_body=original_mail_body,
            profile_text=profile_text,
            conversation_history=conversation_history,
            is_w2_fulltime=is_w2,
            instructions=body.instructions or "",
        )

        subject = result.get('subject', '')
        body_text = result.get('body', '')

        # Save to file in company folder
        try:
            if company_dir and os.path.isdir(company_dir):
                fu_path = os.path.join(company_dir, "follow_up_draft.txt")
                with open(fu_path, "w", encoding="utf-8") as f:
                    f.write(f"Subject: {subject}\n\n{body_text}")
        except Exception as e:
            logger.warning(f"Could not save follow-up file: {e}")

        # Save to DB
        try:
            save_follow_up_draft(record_id, subject, body_text)
        except Exception as e:
            logger.warning(f"Could not save follow-up to DB: {e}")

        response_data = {
            "to_emails": to_emails,
            "subject": subject,
            "body": body_text,
            "w2_detected": is_w2,
        }

        if is_w2:
            try:
                draft_result = gmail_service.save_draft(
                    to_emails=to_emails,
                    subject=result.get('subject', ''),
                    body=result.get('body', ''),
                )
                response_data["auto_draft_saved"] = True
                response_data["draft_id"] = draft_result.get("draft_id", "")
            except Exception as e:
                logger.warning(f"Auto-save draft failed for W2 reply: {e}")
                response_data["auto_draft_saved"] = False

        return response_data
    except RuntimeError as e:
        logger.warning(f"AI provider error: {e}")
        raise HTTPException(status_code=503, detail="The AI provider is unavailable. Please try again.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Follow-up generation error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate follow-up")


@app.post("/api/follow-up/standalone")
@limiter.limit("5/minute")
async def api_standalone_follow_up(request: Request, body: FollowUpRequest):
    """Generate a follow-up reply without a company record — for direct/personal emails."""
    try:
        profile_text = ""
        profile_path = os.path.join(DATA_DIR, "profile.txt")
        if os.path.exists(profile_path):
            with open(profile_path, "r", encoding="utf-8") as f:
                profile_text = f.read().strip()

        to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', body.received_email)))

        is_w2 = detect_w2_fulltime(body.received_email)

        conversation_history = ""
        if to_emails:
            try:
                sender_email = to_emails[0]
                convos = gmail_service.get_conversation_with_sender(sender_email, max_results=20)
                if convos:
                    parts = []
                    for msg in convos:
                        parts.append(
                            f"--- Email ({msg['date']}) ---\n"
                            f"From: {msg['from']}\n"
                            f"To: {msg['to']}\n"
                            f"Subject: {msg['subject']}\n\n"
                            f"{msg['body']}"
                        )
                    conversation_history = "\n\n".join(parts)
            except Exception as e:
                logger.warning(f"Could not fetch conversation history: {e}")

        result = generate_follow_up(
            received_email=body.received_email,
            profile_text=profile_text,
            conversation_history=conversation_history,
            is_w2_fulltime=is_w2,
            instructions=body.instructions or "",
        )

        subject = result.get('subject', '')
        body_text = result.get('body', '')

        # Save to standalone follow-ups folder
        try:
            fu_dir = os.path.join(DATA_DIR, "follow_ups")
            os.makedirs(fu_dir, exist_ok=True)
            from datetime import datetime as _dt
            ts = _dt.now().strftime("%Y%m%d_%H%M%S")
            fu_path = os.path.join(fu_dir, f"follow_up_{ts}.txt")
            with open(fu_path, "w", encoding="utf-8") as f:
                f.write(f"Subject: {subject}\n\n{body_text}")
        except Exception as e:
            logger.warning(f"Could not save standalone follow-up file: {e}")

        response_data = {
            "to_emails": to_emails,
            "subject": subject,
            "body": body_text,
            "w2_detected": is_w2,
        }

        if is_w2:
            try:
                draft_result = gmail_service.save_draft(
                    to_emails=to_emails,
                    subject=subject,
                    body=body_text,
                )
                response_data["auto_draft_saved"] = True
                response_data["draft_id"] = draft_result.get("draft_id", "")
            except Exception as e:
                logger.warning(f"Auto-save draft failed for standalone W2 reply: {e}")
                response_data["auto_draft_saved"] = False

        return response_data
    except RuntimeError as e:
        logger.warning(f"AI provider error: {e}")
        raise HTTPException(status_code=503, detail="The AI provider is unavailable. Please try again.")
    except Exception as e:
        logger.error(f"Standalone follow-up error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate follow-up")

# ─── Job Matcher ───

@app.post("/api/job-matcher/fetch-url")
@limiter.limit("20/minute")
async def fetch_jd_from_url(request: Request, url: str = Form(...)):
    """Fetch job description text from a URL."""
    try:
        if not url or not url.startswith(('http://', 'https://')):
            raise HTTPException(status_code=400, detail="Invalid URL — must start with http:// or https://")
        # Run the blocking HTTP + AI call in a thread pool so it doesn't stall the event loop
        text = await asyncio.get_event_loop().run_in_executor(None, _scrape_jd_from_url, url)
        return {"jd_text": text, "url": url}
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"URL fetch error: {e}", exc_info=True)
        # Surface a readable message for common failures
        msg = str(e)
        if "403" in msg or "999" in msg or "Forbidden" in msg:
            detail = "This site blocked the request (403/bot protection). Paste the JD text directly instead."
        elif "404" in msg:
            detail = "Job posting not found (404). The listing may have expired."
        elif "timeout" in msg.lower() or "timed out" in msg.lower():
            detail = "The request timed out. The site may be slow — try pasting the JD text directly."
        elif "SSL" in msg or "certificate" in msg.lower():
            detail = "SSL error connecting to this site. Try pasting the JD text directly."
        else:
            detail = f"Could not fetch URL: {msg}"
        raise HTTPException(status_code=400, detail=detail)

@app.post("/api/job-matcher/analyze")
@limiter.limit("30/minute")
async def analyze_job(request: Request, jd_text: str = Form(...), source_url: Optional[str] = Form(None)):
    """
    Analyze job description against user profile.
    Saves all results (pass/reject) to history DB and CSV.
    """
    try:
        if not jd_text or len(jd_text) < 50:
            raise HTTPException(status_code=400, detail="Job description too short (minimum 50 characters)")
        if len(jd_text) > 50000:
            raise HTTPException(status_code=400, detail="Job description too long (maximum 50000 characters)")

        # Duplicate JD check
        all_history = get_all_resumes(limit=1000)
        for item in all_history:
            if not item.get('jd_text'):
                continue
            ratio = difflib.SequenceMatcher(None, jd_text, item['jd_text']).ratio()
            if ratio >= 0.95:
                prev = item.get('company_name', 'Unknown')
                raise HTTPException(
                    status_code=400,
                    detail=f"Duplicate JD — already scanned for {prev}. Check your Production Log."
                )

        company_name = _extract_company_name(jd_text)

        # Rule 1: Check lead role (HARD REJECT)
        lead_reason = _check_lead_role(jd_text)
        if lead_reason:
            reason = f"Rejected: {lead_reason}"
            record_id = save_job_matcher_record(company_name, jd_text, 0, False, reason, source_url)
            append_job_matcher_to_csv(company_name, jd_text, 0, "Rejected", reason, source_url or "")
            return {"id": record_id, "error": reason, "can_apply": False, "hard_reject": True, "company_name": company_name}

        # Rule 2: Check visa eligibility (HARD REJECT if excludes GC)
        gc_reason = _check_visa_eligibility(jd_text)
        if gc_reason:
            record_id = save_job_matcher_record(company_name, jd_text, 0, False, gc_reason, source_url)
            append_job_matcher_to_csv(company_name, jd_text, 0, "Rejected", gc_reason, source_url or "")
            return {"id": record_id, "error": gc_reason, "can_apply": False, "hard_reject": True, "company_name": company_name}

        # Rule 3: Check foreign language requirement (HARD REJECT)
        lang_reason = _check_foreign_language(jd_text)
        if lang_reason:
            reason = f"Rejected: {lang_reason}"
            record_id = save_job_matcher_record(company_name, jd_text, 0, False, reason, source_url)
            append_job_matcher_to_csv(company_name, jd_text, 0, "Rejected", reason, source_url or "")
            return {"id": record_id, "error": reason, "can_apply": False, "hard_reject": True, "company_name": company_name}

        # Rule 4: Check employment type (HARD REJECT for W2 only)
        employment_type, emp_warning = _check_employment_type(jd_text)
        if emp_warning:
            reason = f"Rejected: {emp_warning}"
            record_id = save_job_matcher_record(company_name, jd_text, 0, False, reason, source_url)
            append_job_matcher_to_csv(company_name, jd_text, 0, "Rejected", reason, source_url or "")
            return {"id": record_id, "error": reason, "can_apply": False, "hard_reject": True, "company_name": company_name}

        # Extract experience
        years_required = _extract_experience_years(jd_text)

        # Soft warnings
        warnings = []
        if years_required > 10:
            experience_match = max(0, 100 - (years_required - 10) * 5)
            warnings.append(f"Job requires {years_required}+ years, you have ~8-10 years ({experience_match}% experience match)")
        if employment_type == 'unknown':
            warnings.append("Employment type unclear — verify this is C2C or C2H before applying")

        # Extract keywords and metadata
        extracted_keywords = extract_job_keywords(jd_text)
        metadata = analyze_job_metadata(jd_text, extracted_keywords)

        # Calculate overall match
        match_score = calculate_match_score(years_required, extracted_keywords, employment_type)

        # Save to DB and CSV
        scan_data = json.dumps({
            "match_percentage": match_score,
            "warnings": warnings,
            "job_category": metadata.get("primary_role", ""),
            "sub_categories": metadata.get("sub_categories", []),
            "extracted_keywords": extracted_keywords,
        })
        record_id = save_job_matcher_record(company_name, jd_text, match_score, True, None, source_url, scan_data)
        append_job_matcher_to_csv(company_name, jd_text, match_score, "Matched", "", source_url or "")

        return {
            "id": record_id,
            "can_apply": True,
            "hard_reject": False,
            "match_percentage": match_score,
            "warnings": warnings,
            "company_name": company_name,
            "job_category": {
                "name": metadata.get("primary_role", "DevOps Engineer"),
                "sub_categories": metadata.get("sub_categories", []),
                "match_confidence": metadata.get("match_confidence", 0),
            },
            "employment_type": metadata.get("employment_type", employment_type),
            "years_required": years_required,
            "experience_years": f"{years_required}+ years" if years_required > 0 else "Not specified",
            "extracted_keywords": extracted_keywords,
            "location": metadata.get("location", "Not specified"),
            "salary_range": metadata.get("salary_range", "Not specified"),
            "visa_requirements": metadata.get("visa_requirements", "Not specified"),
            "clearance_level": metadata.get("clearance_level", "Not specified"),
            "source_url": source_url,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Job matcher analyze error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to analyze job description")


@app.post("/api/job-matcher/apply")
@limiter.limit("10/minute")
async def apply_to_job(request: Request, jd_text: str = Form(...), resume: Optional[UploadFile] = File(None), selected_resume: Optional[str] = Form(None)):
    """
    Apply to a job by tailoring resume.
    Reuses existing scan logic with job matcher context.
    """
    try:
        # Call the internal function directly, not the @app.post-decorated scan_resume —
        # that wrapper's parameters default to FastAPI's Form(...) marker objects, which
        # only resolve to real values when FastAPI's own request pipeline invokes it.
        # Calling it as a plain function (as this did previously) leaves unset params
        # holding the marker object itself instead of None, breaking any code that
        # touches them (e.g. `rerun_id > 0`).
        result = await _scan_resume_core(request, jd_text, resume, selected_resume, ai_notes="From Job Matcher")
        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Job matcher apply error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to apply to job")

# ─── Telegram Integration ───

def _telegram_process_jd(jd_text: str, chat_id: int):
    """Background task: process a JD received via Telegram and reply with the result."""
    from services.telegram_service import send_message_sync

    try:
        resumes = []
        if os.path.exists("original"):
            for f in os.listdir("original"):
                if f.endswith(".docx"):
                    path = os.path.join("original", f)
                    resumes.append((path, os.path.getmtime(path)))
        if not resumes:
            send_message_sync(chat_id, "No resume found. Please upload a resume via the web dashboard first.")
            return

        resumes.sort(key=lambda x: x[1], reverse=True)
        resume_path = resumes[0][0]
        resume_filename = os.path.basename(resume_path)

        with open(resume_path, "rb") as f:
            file_bytes = f.read()

        if len(jd_text) < 50:
            send_message_sync(chat_id, "JD too short (minimum 50 characters). Please send the full job description.")
            return

        years = _extract_experience_years(jd_text)
        if years > 10:
            send_message_sync(chat_id, f"Skipped — requires {years}+ years experience (max: 10)")
            return

        gc_reason = _check_visa_eligibility(jd_text)
        if gc_reason:
            send_message_sync(chat_id, f"Skipped — {gc_reason}")
            return

        lang_reason = _check_foreign_language(jd_text)
        if lang_reason:
            send_message_sync(chat_id, f"Skipped — {lang_reason}")
            return

        lead_reason = _check_lead_role(jd_text)
        if lead_reason:
            send_message_sync(chat_id, f"Skipped — {lead_reason}")
            return

        result = _process_single_jd(jd_text, file_bytes, resume_path)

        if result.get("skipped"):
            send_message_sync(chat_id, f"Skipped — {result.get('reason', 'Unknown reason')}")
            return

        company = result.get("company_name", "Unknown")
        score = result.get("score", 0)
        after_score = result.get("after_score", score)
        tailored = result.get("tailored", False)
        missing = result.get("missing_keywords", [])

        record_id = result.get("id", 0)

        lines = [f"Processed JD for {company}!"]
        if tailored and after_score != score:
            lines.append(f"Score: {score}% -> {after_score}% (+{after_score - score})")
        else:
            lines.append(f"Match Score: {after_score}%")

        if tailored:
            lines.append("Resume has been tailored")
        else:
            lines.append("Resume already strong — no changes needed")

        if missing:
            lines.append(f"Missing keywords: {', '.join(missing[:8])}")

        lines.append(f"\nUsing resume: {resume_filename}")
        lines.append("Files saved — check your dashboard for downloads.")

        buttons = {
            "inline_keyboard": [
                [
                    {"text": "Generate Cover Letter", "callback_data": f"cover_letter:{record_id}"},
                    {"text": "Generate Mail Draft", "callback_data": f"mail_draft:{record_id}"},
                ],
                [
                    {"text": "Save Draft to Gmail", "callback_data": f"gmail_draft:{record_id}"},
                ],
            ]
        }

        send_message_sync(chat_id, "\n".join(lines), reply_markup=buttons)
        logger.info(f"Telegram JD processed: {company} (score={after_score}%) chat={chat_id}")

    except Exception as e:
        logger.error(f"Telegram processing error: {e}", exc_info=True)
        try:
            send_message_sync(chat_id, f"Processing failed: {str(e)[:300]}")
        except Exception:
            logger.error("Failed to send Telegram error reply")


def _telegram_generate_cover_letter(record_id: int, chat_id: int):
    """Background task: generate cover letter and send result via Telegram."""
    from services.telegram_service import send_message_sync
    try:
        record = get_resume_by_id(record_id)
        if not record:
            send_message_sync(chat_id, "Record not found.")
            return

        file_path = _resolve_resume_path(record)
        if not file_path:
            send_message_sync(chat_id, "Resume file not found for this record.")
            return

        with open(file_path, "rb") as f:
            resume_text = extract_text_from_docx(f.read())

        cover_letter = generate_cover_letter(resume_text, record['jd_text'], record['company_name'])

        company_dir = os.path.dirname(os.path.abspath(file_path))
        cl_filename = get_output_filename("cover_letter", record_id)
        cl_path = os.path.join(company_dir, cl_filename).replace("\\", "/")

        cl_doc = docx.Document()
        cl_doc.add_paragraph(cover_letter)
        cl_doc.save(cl_path)

        company = record.get('company_name', 'Unknown')
        msg = f"Cover letter generated for {company}!\n\n{cover_letter[:3500]}"
        buttons = {
            "inline_keyboard": [
                [
                    {"text": "Generate Mail Draft", "callback_data": f"mail_draft:{record_id}"},
                    {"text": "Save Draft to Gmail", "callback_data": f"gmail_draft:{record_id}"},
                ],
            ]
        }
        send_message_sync(chat_id, msg, reply_markup=buttons)
        logger.info(f"Telegram cover letter generated for record {record_id}")

    except Exception as e:
        logger.error(f"Telegram cover letter error: {e}", exc_info=True)
        try:
            send_message_sync(chat_id, f"Cover letter generation failed: {str(e)[:300]}")
        except Exception:
            pass


def _telegram_generate_mail_draft(record_id: int, chat_id: int):
    """Background task: generate mail draft and send result via Telegram."""
    from services.telegram_service import send_message_sync
    try:
        record = get_resume_by_id(record_id)
        if not record:
            send_message_sync(chat_id, "Record not found.")
            return

        file_path = _resolve_resume_path(record) or record.get('file_path')
        if not file_path:
            send_message_sync(chat_id, "No file path for this record.")
            return

        company_dir = os.path.dirname(file_path)
        resume_text = ""
        cover_letter_text = ""

        if os.path.exists(file_path):
            with open(file_path, "rb") as f:
                resume_text = extract_text_from_docx(f.read())

        if os.path.isdir(company_dir):
            for filename in os.listdir(company_dir):
                if ("cover" in filename.lower() or filename.startswith("cover_letter_")) and filename.endswith(".docx"):
                    filepath = os.path.join(company_dir, filename)
                    if filepath != file_path:
                        with open(filepath, "rb") as f:
                            cover_letter_text = extract_text_from_docx(f.read())
                        break

        if not resume_text:
            send_message_sync(chat_id, "Could not read resume text.")
            return

        jd_text = record.get('jd_text', '')
        if not jd_text:
            send_message_sync(chat_id, "No job description found for this record.")
            return

        profile_text = ""
        profile_path = os.path.join(DATA_DIR, "profile.txt")
        if os.path.exists(profile_path):
            with open(profile_path, "r", encoding="utf-8") as f:
                profile_text = f.read().strip()

        mail = generate_mail_draft(resume_text, jd_text, cover_letter_text, record['company_name'], profile_text)
        subject = mail.get('subject', '')
        body = mail.get('body', '')

        draft_filename = get_output_filename("mail_draft", record_id)
        draft_path = os.path.join(company_dir, draft_filename).replace("\\", "/")
        with open(draft_path, "w", encoding="utf-8") as f:
            f.write(f"Subject: {subject}\n\n{body}")

        company = record.get('company_name', 'Unknown')
        to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', jd_text)))
        to_line = f"To: {', '.join(to_emails)}" if to_emails else "To: (no email found in JD)"

        msg = f"Mail draft for {company}:\n\n{to_line}\nSubject: {subject}\n\n{body[:3000]}"
        buttons = {
            "inline_keyboard": [
                [{"text": "Save Draft to Gmail", "callback_data": f"gmail_draft:{record_id}"}],
            ]
        }
        send_message_sync(chat_id, msg, reply_markup=buttons)
        logger.info(f"Telegram mail draft generated for record {record_id}")

    except Exception as e:
        logger.error(f"Telegram mail draft error: {e}", exc_info=True)
        try:
            send_message_sync(chat_id, f"Mail draft generation failed: {str(e)[:300]}")
        except Exception:
            pass


def _telegram_save_gmail_draft(record_id: int, chat_id: int):
    """Background task: generate mail + save as Gmail draft via Telegram."""
    from services.telegram_service import send_message_sync
    try:
        record = get_resume_by_id(record_id)
        if not record:
            send_message_sync(chat_id, "Record not found.")
            return

        file_path = _resolve_resume_path(record) or record.get('file_path')
        if not file_path:
            send_message_sync(chat_id, "No file path for this record.")
            return

        company_dir = os.path.dirname(file_path)
        jd_text = record.get('jd_text', '')
        company = record.get('company_name', 'Unknown')

        subject = ""
        body = ""
        draft_files = [f for f in os.listdir(company_dir) if f.startswith("mail_draft_") and f.endswith(".txt")] if os.path.isdir(company_dir) else []
        if draft_files:
            with open(os.path.join(company_dir, draft_files[0]), "r", encoding="utf-8") as f:
                content = f.read()
            if content.startswith("Subject:"):
                sep = content.find('\n\n')
                subject = content[len("Subject:"):sep].strip() if sep != -1 else content[len("Subject:"):].strip()
                body = content[sep + 2:] if sep != -1 else ""
            else:
                body = content

        if not body:
            resume_text = ""
            cover_letter_text = ""
            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    resume_text = extract_text_from_docx(f.read())

            if os.path.isdir(company_dir):
                for filename in os.listdir(company_dir):
                    if ("cover" in filename.lower() or filename.startswith("cover_letter_")) and filename.endswith(".docx"):
                        filepath = os.path.join(company_dir, filename)
                        if filepath != file_path:
                            with open(filepath, "rb") as f:
                                cover_letter_text = extract_text_from_docx(f.read())
                            break

            profile_text = ""
            profile_path = os.path.join(DATA_DIR, "profile.txt")
            if os.path.exists(profile_path):
                with open(profile_path, "r", encoding="utf-8") as f:
                    profile_text = f.read().strip()

            mail = generate_mail_draft(resume_text, jd_text, cover_letter_text, company, profile_text)
            subject = mail.get('subject', '')
            body = mail.get('body', '')

            draft_filename = get_output_filename("mail_draft", record_id)
            draft_path = os.path.join(company_dir, draft_filename).replace("\\", "/")
            with open(draft_path, "w", encoding="utf-8") as f:
                f.write(f"Subject: {subject}\n\n{body}")

        to_emails = list(set(re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', jd_text)))

        result = gmail_service.save_draft(
            to_emails=to_emails,
            subject=subject,
            body=body,
            attachment_path=file_path if os.path.exists(file_path) else None,
        )

        attached = result.get('attachments', [])
        to_line = ', '.join(to_emails) if to_emails else '(no email found)'
        msg = f"Gmail draft saved for {company}!\n\nTo: {to_line}\nSubject: {subject}\nAttachments: {', '.join(attached) if attached else 'none'}\n\nCheck your Gmail Drafts folder."
        send_message_sync(chat_id, msg)
        logger.info(f"Telegram Gmail draft saved for record {record_id}")

    except RuntimeError as e:
        error_msg = str(e)
        if "Gmail not connected" in error_msg:
            send_message_sync(chat_id, "Gmail not connected. Please connect your Gmail account via the web dashboard first.")
        else:
            send_message_sync(chat_id, f"Gmail draft failed: {error_msg[:300]}")
    except Exception as e:
        logger.error(f"Telegram Gmail draft error: {e}", exc_info=True)
        try:
            send_message_sync(chat_id, f"Gmail draft failed: {str(e)[:300]}")
        except Exception:
            pass


def _telegram_truncate(text: str, max_len: int = 50) -> str:
    """Some legacy records store the FULL scraped job text in the 'title' field
    (jd_text), not a short title — without this, a single item can turn into a
    multi-KB wall of text in a Telegram message."""
    text = (text or '').strip().split('\n')[0]
    return text[:max_len].strip() + '...' if len(text) > max_len else text


def _telegram_matches_text(limit: int = 5) -> str:
    from database import get_found_jobs
    jobs = get_found_jobs(limit=limit, offset=0)
    if not jobs:
        return "No Command Center matches yet. Run a search from the app first."
    lines = ["Recent matches:"]
    for j in jobs:
        lines.append(f"- {j.get('company', '')} - {_telegram_truncate(j.get('title', ''))} ({j.get('score', 0)}%)")
    return "\n".join(lines)


def _telegram_followups_text() -> str:
    from database import get_action_queue
    queue = get_action_queue(item_limit=10)
    items = queue.get('follow_ups_due', {}).get('items', [])
    if not items:
        return "No follow-ups due right now."
    lines = ["Follow-ups due:"]
    for item in items:
        days = item.get('days_since')
        lines.append(f"- {item.get('company', '')} - {_telegram_truncate(item.get('title', ''))}" + (f" ({days}d)" if days is not None else ""))
    return "\n".join(lines)


def _telegram_queue_text() -> str:
    from database import get_action_queue
    from services.telegram_notifier import SECTION_LABELS
    queue = get_action_queue()
    lines = ["Action Queue:"]
    for key, label in SECTION_LABELS.items():
        count = queue.get(key, {}).get('count', 0)
        lines.append(f"- {label}: {count}")
    return "\n".join(lines)


_telegram_polling = False
_telegram_msg_buffers: dict[int, list[str]] = {}
_telegram_flush_tasks: dict[int, asyncio.Task] = {}
_TELEGRAM_BUFFER_DELAY = 3.0

async def _telegram_flush_buffer(chat_id: int):
    """Wait for the buffer delay, then combine and process all buffered messages for a chat.
    Splits on '---' separator to handle multiple JDs in one stream."""
    await asyncio.sleep(_TELEGRAM_BUFFER_DELAY)

    from services import telegram_service

    parts = _telegram_msg_buffers.pop(chat_id, [])
    _telegram_flush_tasks.pop(chat_id, None)

    if not parts:
        return

    combined_text = "\n".join(parts)

    jds = re.split(r'\n\s*---\s*\n', combined_text)
    jds = [jd.strip() for jd in jds if jd.strip()]

    if not jds:
        return

    logger.info(f"Telegram buffer flushed for chat {chat_id}: {len(parts)} part(s), {len(jds)} JD(s), {len(combined_text)} chars total")

    loop = asyncio.get_event_loop()
    if len(jds) == 1:
        await telegram_service.send_message(chat_id, "Processing your JD... You'll get a reply shortly.")
        loop.run_in_executor(None, _telegram_process_jd, jds[0], chat_id)
    else:
        await telegram_service.send_message(chat_id, f"Found {len(jds)} JDs (separated by ---). Processing each one...")
        for i, jd in enumerate(jds):
            loop.run_in_executor(None, _telegram_process_jd, jd, chat_id)

async def _telegram_poll_loop():
    """Long-polling loop that listens for Telegram messages."""
    global _telegram_polling
    from services import telegram_service

    if not telegram_service.is_configured():
        logger.info("Telegram bot not configured — polling disabled")
        return

    await asyncio.sleep(2)

    try:
        bot_info = await telegram_service.get_me()
        logger.info(f"Telegram bot connected: @{bot_info.get('username', '?')}")
    except Exception as e:
        logger.error(f"Telegram bot connection failed: {e}")
        return

    try:
        import httpx as _hx
        async with _hx.AsyncClient(timeout=10) as client:
            await client.post(f"https://api.telegram.org/bot{os.getenv('TELEGRAM_BOT_TOKEN')}/deleteWebhook")
    except Exception:
        pass

    _telegram_polling = True
    offset = 0

    while _telegram_polling:
        try:
            updates = await telegram_service.get_updates(offset=offset, timeout=30)
            for update in updates:
                offset = update["update_id"] + 1

                callback = update.get("callback_query")
                if callback:
                    cb_id = callback.get("id")
                    cb_data = callback.get("data", "")
                    cb_chat_id = callback.get("message", {}).get("chat", {}).get("id")
                    if cb_chat_id and ":" in cb_data:
                        action, rec_id_str = cb_data.split(":", 1)
                        try:
                            rec_id = int(rec_id_str)
                        except ValueError:
                            await telegram_service.answer_callback_query(cb_id, "Invalid record.")
                            continue

                        loop = asyncio.get_event_loop()
                        if action == "cover_letter":
                            await telegram_service.answer_callback_query(cb_id, "Generating cover letter...")
                            await telegram_service.send_message(cb_chat_id, "Generating cover letter... please wait.")
                            loop.run_in_executor(None, _telegram_generate_cover_letter, rec_id, cb_chat_id)
                        elif action == "mail_draft":
                            await telegram_service.answer_callback_query(cb_id, "Generating mail draft...")
                            await telegram_service.send_message(cb_chat_id, "Generating mail draft... please wait.")
                            loop.run_in_executor(None, _telegram_generate_mail_draft, rec_id, cb_chat_id)
                        elif action == "gmail_draft":
                            await telegram_service.answer_callback_query(cb_id, "Saving draft to Gmail...")
                            await telegram_service.send_message(cb_chat_id, "Generating mail and saving to Gmail drafts... please wait.")
                            loop.run_in_executor(None, _telegram_save_gmail_draft, rec_id, cb_chat_id)
                        else:
                            await telegram_service.answer_callback_query(cb_id, "Unknown action.")
                    else:
                        await telegram_service.answer_callback_query(cb_id)
                    continue

                msg = update.get("message", {})
                text = msg.get("text", "").strip()
                chat_id = msg.get("chat", {}).get("id")

                if not text or not chat_id:
                    continue

                # Remember this chat so proactive pushes (new matches, daily digest,
                # periodic action-queue check) have somewhere to send to — Telegram has
                # no "message whoever set up the bot" API, only reply-to-an-existing-chat.
                telegram_service.add_known_chat_id(chat_id)

                if text.startswith("/matches"):
                    await telegram_service.send_message(chat_id, _telegram_matches_text())
                    continue

                if text.startswith("/followups"):
                    await telegram_service.send_message(chat_id, _telegram_followups_text())
                    continue

                if text.startswith("/queue"):
                    await telegram_service.send_message(chat_id, _telegram_queue_text())
                    continue

                if text.startswith("/start"):
                    await telegram_service.send_message(chat_id, "Welcome to Job Tailored Resume Bot!\n\nSend me a Job Description and I'll:\n- Analyze it against your resume\n- Tailor your resume automatically\n- Save everything to your dashboard\n\nJust paste a JD and hit send.\n\nCommand Center:\n- /matches - recent auto-search matches\n- /followups - applications due for a follow-up\n- /queue - Action Queue summary\nYou'll also get pushed here automatically when new matches or action items show up.\n\nTips:\n- Long JDs split across messages are auto-combined\n- Use --- between JDs to send multiple at once\n- Use /scan to start a new JD (flushes any pending text)")
                    continue

                if text.startswith("/status"):
                    resumes = [f for f in os.listdir("original") if f.endswith(".docx")] if os.path.exists("original") else []
                    await telegram_service.send_message(chat_id, f"Bot is running.\nResumes available: {len(resumes)}")
                    continue

                if text.startswith("/scan"):
                    existing_task = _telegram_flush_tasks.get(chat_id)
                    if existing_task and not existing_task.done():
                        existing_task.cancel()
                    buffered = _telegram_msg_buffers.pop(chat_id, [])
                    if buffered:
                        combined = "\n".join(buffered)
                        logger.info(f"Telegram /scan flushed previous buffer for chat {chat_id}: {len(combined)} chars")
                        await telegram_service.send_message(chat_id, "Processing previous JD before starting new one...")
                        loop = asyncio.get_event_loop()
                        loop.run_in_executor(None, _telegram_process_jd, combined, chat_id)
                    jd_after_command = text[len("/scan"):].strip()
                    if jd_after_command:
                        _telegram_msg_buffers[chat_id] = [jd_after_command]
                        _telegram_flush_tasks[chat_id] = asyncio.create_task(_telegram_flush_buffer(chat_id))
                    else:
                        await telegram_service.send_message(chat_id, "Ready for new JD. Paste it now.")
                    continue

                if text.startswith("/"):
                    await telegram_service.send_message(chat_id, "Commands:\n/start - Welcome message\n/status - Check bot status\n/scan - Start a new JD (flushes pending text)\n/matches - Recent Command Center matches\n/followups - Applications due for a follow-up\n/queue - Action Queue summary\n\nOr just send a Job Description to process it.\nUse --- between JDs to send multiple at once.")
                    continue

                logger.info(f"Telegram message from chat {chat_id}: {len(text)} chars")

                if chat_id not in _telegram_msg_buffers:
                    _telegram_msg_buffers[chat_id] = []
                _telegram_msg_buffers[chat_id].append(text)

                existing_task = _telegram_flush_tasks.get(chat_id)
                if existing_task and not existing_task.done():
                    existing_task.cancel()

                _telegram_flush_tasks[chat_id] = asyncio.create_task(_telegram_flush_buffer(chat_id))

        except asyncio.CancelledError:
            _telegram_polling = False
            break
        except Exception as e:
            if str(e):
                logger.error(f"Telegram polling error: {e}")
            await asyncio.sleep(5)


@app.on_event("startup")
async def start_telegram_polling():
    from services import telegram_service
    if telegram_service.is_configured():
        asyncio.create_task(_telegram_poll_loop())


@app.on_event("startup")
async def start_telegram_notify_loop():
    from services import telegram_service, telegram_notifier
    if telegram_service.is_configured():
        asyncio.create_task(telegram_notifier.telegram_notify_loop())


@app.on_event("startup")
async def start_daily_search_scheduler():
    from services import scheduler as scheduler_service
    asyncio.create_task(scheduler_service.daily_search_loop(_run_command_center_search))


@app.on_event("startup")
async def start_inbox_reply_check_loop():
    from services import inbox_matcher
    asyncio.create_task(inbox_matcher.inbox_reply_check_loop())


@app.get("/api/telegram/status")
async def telegram_status():
    """Check if Telegram bot is configured and running."""
    from services import telegram_service
    result = {"configured": telegram_service.is_configured(), "polling": _telegram_polling}
    if telegram_service.is_configured():
        try:
            bot_info = await telegram_service.get_me()
            result["bot_username"] = bot_info.get("username", "")
        except Exception:
            result["bot_username"] = ""
    return result



PIPELINE_MATCH_SCORE_THRESHOLD = 70

@app.get("/api/command-center/dashboard")
@limiter.limit("60/minute")
async def get_command_center_dashboard(request: Request):
    try:
        from database import get_all_resumes, get_found_jobs, get_action_queue, get_job_source_breakdown, get_skipped_jobs, get_career_intel
        records = get_all_resumes(1000, 0)

        # 7-stage pipeline: Discovered/Matched split the raw 'Found' bucket by score
        # (Matched = scored >= threshold, i.e. a real hit vs. just a scraped listing).
        # Discovered/Matched are scoped to command-center auto-search results only —
        # plain resume-tailoring scans from the Dashboard ('Scanned' status) aren't part
        # of the job-discovery funnel and would otherwise inflate these buckets. Later
        # stages (Saved/Applied/Interview/Offer/Rejected) apply regardless of source,
        # since a job can be moved through the funnel from any entry point.
        pipeline = {'Discovered': 0, 'Matched': 0, 'Saved': 0, 'Applied': 0, 'Interview': 0, 'Offer': 0, 'Rejected': 0}

        for r in records:
            status = r.get('status', 'Found')
            score = r.get('score') or 0
            if status == 'Found' and r.get('source') == 'command-center':
                if score >= PIPELINE_MATCH_SCORE_THRESHOLD:
                    pipeline['Matched'] += 1
                else:
                    pipeline['Discovered'] += 1
            elif status in ('Shortlisted', 'Matched'):
                pipeline['Saved'] += 1
            elif status == 'Applied':
                pipeline['Applied'] += 1
            elif status == 'Interviewing':
                pipeline['Interview'] += 1
            elif status == 'Offered':
                pipeline['Offer'] += 1
            elif status == 'Rejected':
                pipeline['Rejected'] += 1

        action_queue = get_action_queue()
        try:
            from services.inbox_matcher import get_unhandled_inbox_replies
            reply_items, reply_count = get_unhandled_inbox_replies()
            action_queue['inbox_replies'] = {'count': reply_count, 'items': reply_items}
        except Exception as e:
            logger.warning(f"Failed to load inbox replies for dashboard: {e}")
            action_queue['inbox_replies'] = {'count': 0, 'items': []}
        top_jobs = [_annotate_employment_type(j) for j in get_found_jobs(5)]

        try:
            from services.scan_status import get_last_scan
            last_scan = get_last_scan()
        except Exception:
            last_scan = None

        try:
            from services.telegram_service import is_configured as telegram_configured
            telegram_connected = telegram_configured() and _telegram_polling
        except Exception:
            telegram_connected = False

        from services.scheduler import get_schedule_info
        schedule_info = get_schedule_info()

        return {
            'metrics': {
                'new_jobs': pipeline['Discovered'] + pipeline['Matched'],
                'strong_matches': pipeline['Matched'],
                'jobs_skip': pipeline['Rejected'],
                'apps_pending': pipeline['Applied'],
                'follow_ups': action_queue['follow_ups_due']['count']
            },
            'pipeline': pipeline,
            'top_jobs': top_jobs,
            'action_queue': action_queue,
            'skipped_jobs': get_skipped_jobs(),
            'job_source_health': get_job_source_breakdown(),
            'last_scan': last_scan,
            'automation': {
                'automation_enabled': schedule_info['enabled'],
                'daily_search_scheduled': schedule_info['enabled'],
                'daily_search_time': f"{schedule_info['hour']:02d}:{schedule_info['minute']:02d} {schedule_info['timezone']}",
                'daily_search_next_run': schedule_info['next_run_at'],
                'daily_search_last_run': schedule_info['last_run_at'],
                'telegram_connected': telegram_connected,
            },
            'career_intel': get_career_intel(),
        }
    except Exception as e:
        logger.error(f"Dashboard error: {e}")
        raise HTTPException(status_code=500, detail="Failed to load dashboard")


@app.post("/api/command-center/inbox-replies/check")
@limiter.limit("10/minute")
async def check_inbox_replies_now(request: Request):
    """Manually trigger an inbox check instead of waiting for the 30-min background
    loop — lets the dashboard feel responsive right after replying to a company."""
    from services.inbox_matcher import check_inbox_for_replies
    new_matches = check_inbox_for_replies()
    return {"new_matches": len(new_matches)}


@app.post("/api/command-center/inbox-replies/{message_id}/dismiss")
@limiter.limit("60/minute")
async def dismiss_inbox_reply(request: Request, message_id: str):
    from services.inbox_matcher import mark_inbox_reply_handled
    found = mark_inbox_reply_handled(message_id)
    if not found:
        raise HTTPException(status_code=404, detail="Inbox reply not found")
    return {"status": "dismissed"}


class ManualJobRequest(BaseModel):
    company: str
    title: str
    url: str = ""
    description: str = ""
    notes: str = ""

@app.post("/api/jobs/manual-add")
@limiter.limit("30/minute")
async def add_manual_job(request: Request, payload: ManualJobRequest):
    if not payload.company.strip() or not payload.title.strip():
        raise HTTPException(status_code=400, detail="Company and title are required")
    from database import save_manual_job
    record_id = save_manual_job({
        'company': payload.company.strip(),
        'title': payload.title.strip(),
        'url': payload.url.strip(),
        'description': payload.description.strip(),
        'notes': payload.notes.strip(),
    })
    return {"id": record_id, "status": "saved"}


@app.get("/api/jobs/matches")
@limiter.limit("60/minute")
async def get_job_matches(request: Request, limit: int = 20, offset: int = 0):
    """All found jobs, best score first, paginated — backs the Command Center's
    'View All Matches' page."""
    from database import get_found_jobs, get_found_jobs_count
    limit = max(1, min(limit, 100))
    offset = max(0, offset)
    jobs = [_annotate_employment_type(j) for j in get_found_jobs(limit=limit, offset=offset)]
    total = get_found_jobs_count()
    return {"jobs": jobs, "total": total, "limit": limit, "offset": offset}


@app.post("/api/jobs/clear")
@limiter.limit("5/minute")
async def clear_command_center_jobs_endpoint(request: Request):
    """Wipe every job the Command Center tracks (auto-search results + manually added
    jobs) so the user can start fresh. Does not touch Resume Tailor or Job Finder data."""
    from database import clear_command_center_jobs
    deleted = clear_command_center_jobs()
    try:
        scan_file = os.path.join(DATA_DIR, "last_scan.json")
        if os.path.exists(scan_file):
            os.remove(scan_file)
    except Exception as e:
        logger.warning(f"Failed to clear last scan status: {e}")
    return {"status": "ok", "deleted": deleted}


class ScheduleUpdate(BaseModel):
    enabled: bool = True
    hour: int = 10
    minute: int = 0

@app.post("/api/automation/schedule")
@limiter.limit("10/minute")
async def update_schedule_endpoint(request: Request, payload: ScheduleUpdate):
    if not (0 <= payload.hour <= 23) or not (0 <= payload.minute <= 59):
        raise HTTPException(status_code=400, detail="hour must be 0-23 and minute 0-59")
    from services.scheduler import set_schedule, get_schedule_info
    set_schedule(payload.enabled, payload.hour, payload.minute)
    return get_schedule_info()


@app.get("/api/automation/schedule")
@limiter.limit("30/minute")
async def get_schedule_endpoint(request: Request):
    from services.scheduler import get_schedule_info
    return get_schedule_info()


EMPLOYMENT_TYPE_LABELS = {
    'c2c': 'C2C', 'c2h': 'C2H', 'w2': 'W2', 'contract': 'Contract', 'unknown': 'Unspecified',
}

def _annotate_employment_type(job: dict) -> dict:
    """Adds employment_type/employment_type_label to a job dict if not already present.
    Command Center auto-search jobs already have this persisted from search time (see
    _score_jobs_with_claude), detected against the FULL scraped posting — the exact text
    the hard-reject rule evaluated. This only recomputes as a fallback for older/manual
    records that predate that (from just the saved `description` field, a narrower slice
    of the original page that can under-detect a contract-type signal sitting outside
    the specific "Description:" section). Mutates and returns job for easy use in a list
    comprehension/loop."""
    if job.get('employment_type_label'):
        return job
    description = (job.get('description') or '').strip()
    if description:
        employment_type, _ = _check_employment_type(description)
        job['employment_type'] = employment_type
        job['employment_type_label'] = EMPLOYMENT_TYPE_LABELS.get(employment_type, 'Unspecified')
    else:
        job['employment_type'] = None
        job['employment_type_label'] = None
    return job

@app.get("/api/jobs/{job_id}")
@limiter.limit("60/minute")
async def get_job_detail_endpoint(request: Request, job_id: int):
    """Full job object for the Command Center's Job Detail Workspace."""
    from database import get_job_detail
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return _annotate_employment_type(job)


# Pipeline statuses as used throughout Command Center (pipeline overview, action
# queue, skip reasons) — kept separate from /api/history/{id}/status's own enum
# ("Scanned"/"Phone Screen"/"Interview"/"Offer") since that endpoint serves the
# Resume Tailor production log and uses different status names for the same stages.
JOB_STATUS_VALUES = ["Found", "Shortlisted", "Applied", "Interviewing", "Offered", "Rejected"]

class JobStatusUpdate(BaseModel):
    status: str
    rejection_reason: Optional[str] = None

@app.post("/api/jobs/{job_id}/status")
@limiter.limit("30/minute")
async def update_job_status_endpoint(request: Request, job_id: int, payload: JobStatusUpdate):
    if payload.status not in JOB_STATUS_VALUES:
        raise HTTPException(status_code=400, detail=f"status must be one of {JOB_STATUS_VALUES}")
    from database import get_resume_by_id, update_resume_status, get_job_detail
    existing = get_resume_by_id(job_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Job not found")
    update_resume_status(job_id, payload.status, rejection_reason=payload.rejection_reason)
    updated_job = get_job_detail(job_id)

    # Each status-transition button (Mark Applied/Interview/Rejected, Skip Job) gets its
    # own artifact file, keyed by the exact status+reason combo the frontend sends —
    # "Rejected by employer" vs "Skipped by user" both map to status='Rejected' but are
    # distinct user actions, so they get distinct files.
    status_filenames = {
        ("Applied", None): "mark_applied.txt",
        ("Interviewing", None): "mark_interview.txt",
        ("Interviewing", "Assessment stage"): "mark_assessment.txt",
        ("Offered", None): "mark_offer.txt",
        ("Rejected", "Rejected by employer"): "mark_rejected.txt",
        ("Rejected", "Skipped by user"): "skip_job.txt",
    }
    filename = status_filenames.get((payload.status, payload.rejection_reason))
    if filename:
        content = (
            f"Status changed to {payload.status} at {datetime.now().isoformat()}\n"
            f"Job: {updated_job.get('title', '')} at {updated_job.get('company', '')}\n"
        )
        if payload.rejection_reason:
            content += f"Reason: {payload.rejection_reason}\n"
        _write_job_artifact(job_id, updated_job.get('company', ''), filename, content)

    return updated_job


@app.post("/api/jobs/{job_id}/rescore")
@limiter.limit("10/minute")
async def rescore_job_endpoint(request: Request, job_id: int):
    """Re-run the same Claude scoring pass used by auto-search against this job's
    already-stored description — for when the user wants a fresh AI opinion."""
    from database import get_job_detail, update_job_score
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    description = (job.get('description') or '').strip()
    if len(description) < 40:
        raise HTTPException(status_code=400, detail="Not enough job description text to re-score. Try 'Fetch job details' first.")

    job_text = f"Title: {job.get('title', '')}\nCompany: {job.get('company', '')}\nLocation: {job.get('location', '')}\nDescription: {description}"
    try:
        results = _score_jobs_with_claude([job_text], job.get('title') or 'this role')
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not results:
        raise HTTPException(status_code=502, detail="AI scoring returned no result — try again.")

    scored = results[0]
    return update_job_score(
        job_id,
        score=scored.get('score', job.get('score', 0)),
        tags=scored.get('tags', []),
        reasons=scored.get('reasons', []),
        next_action=scored.get('next_action', ''),
    )


class DescriptionUpdate(BaseModel):
    description: str

@app.post("/api/jobs/{job_id}/description")
@limiter.limit("30/minute")
async def paste_job_description_endpoint(request: Request, job_id: int, payload: DescriptionUpdate):
    if len(payload.description.strip()) < 20:
        raise HTTPException(status_code=400, detail="Description is too short to be useful")
    from database import get_resume_by_id, update_job_description
    if not get_resume_by_id(job_id):
        raise HTTPException(status_code=404, detail="Job not found")
    return update_job_description(job_id, payload.description.strip())


@app.post("/api/jobs/{job_id}/fetch-description")
@limiter.limit("10/minute")
async def fetch_job_description_endpoint(request: Request, job_id: int):
    """Re-scrape the posting URL for full description text — same fetch approach as
    the auto-search scraping fallback, just targeted at one already-saved job."""
    from database import get_job_detail, update_job_description
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    url = job.get('url')
    if not url:
        raise HTTPException(status_code=400, detail="This job has no source URL to fetch from")
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        resp = requests.get(url, headers=headers, timeout=10)
        if resp.status_code != 200:
            raise HTTPException(status_code=502, detail=f"Fetch failed with status {resp.status_code}")
        soup = BeautifulSoup(resp.text, 'html.parser')
        # Match the cap used everywhere else JDs are scraped (auto-search, contact
        # discovery) — this was capped lower here, so a manual "Fetch description"
        # re-fetch could produce a shorter JD than the one auto-search originally found.
        text = soup.get_text(separator=' ', strip=True)[:20000]
        if len(text.strip()) < 40:
            raise HTTPException(status_code=502, detail="Fetched page had no usable text")
        return update_job_description(job_id, text.strip())
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"Failed to fetch description for job {job_id}: {e}")
        raise HTTPException(status_code=502, detail="Could not fetch the job page — try pasting the description manually")


@app.post("/api/jobs/{job_id}/tailor")
@limiter.limit("10/minute")
async def tailor_job_endpoint(request: Request, job_id: int, selected_resume: Optional[str] = None):
    """Tailor a resume against this job's own description without leaving the Command
    Center workspace. Reuses the existing /api/scan flow directly (same function call,
    not a re-implementation) so there's exactly one resume-tailoring code path in the
    app; this just calls it with the job's description as the JD and links the
    resulting file back onto this job so Generate Cover Letter etc. can find it."""
    from database import get_job_detail, link_tailored_resume
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    jd_text = (job.get('description') or '').strip()
    if len(jd_text) < 50:
        raise HTTPException(status_code=400, detail="Not enough job description text to tailor against — fetch or paste the description first.")

    job_dir = _job_artifact_dir(job_id, job.get('company') or 'Unknown')
    scan_response = await _scan_resume_core(
        request=request,
        jd_text=jd_text,
        resume=None,
        selected_resume=selected_resume,
        ai_notes=None,
        rerun_id=None,
        storage_root="online-platform",
        override_company_dir=job_dir,
        # Re-tailoring the same Command Center job produces the same JD text every
        # time by design (it's this job's own stored description) — the duplicate-JD
        # check exists to catch accidental double-submissions elsewhere, not this.
        # Without this flag, clicking "Generate Tailored Resume" a second time for the
        # same job would 409 against the scan record the first click created.
        skip_duplicate_check=True,
    )

    file_path = scan_response.get('file_path') if isinstance(scan_response, dict) else None
    updated_job = link_tailored_resume(job_id, file_path) if file_path else job
    return {"tailor_result": scan_response, "job": updated_job}


@app.post("/api/jobs/{job_id}/explain-match")
@limiter.limit("30/minute")
async def explain_match_endpoint(request: Request, job_id: int):
    """Persist the AI match explanation (score/reasons/tags/next_action already produced
    by scoring) as a standalone artifact — doesn't call the AI again, just writes down
    what it already said the last time this job was scored."""
    from database import get_job_detail
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if not job.get('reasons'):
        raise HTTPException(status_code=400, detail="This job hasn't been scored yet — score it first.")

    lines = [
        f"Match Explanation — {job.get('title', '')} at {job.get('company', '')}",
        f"Score: {job.get('score', 0)}",
        "",
        "Reasons:",
    ]
    lines += [f"- {r}" for r in job.get('reasons', [])]
    if job.get('tags'):
        lines += ["", f"Tags: {', '.join(job['tags'])}"]
    if job.get('next_action'):
        lines += [f"Suggested next action: {job['next_action']}"]
    content = "\n".join(lines) + "\n"

    _write_job_artifact(job_id, job.get('company', ''), "explain_match.txt", content)
    return {"explain_match": content}


@app.post("/api/jobs/{job_id}/send-to-tailor")
@limiter.limit("30/minute")
async def send_to_tailor_log_endpoint(request: Request, job_id: int):
    """Log that this job's JD was handed off to the Dashboard's standalone resume
    tailor flow — that flow isn't job-linked (it writes to trailerd/, not this job's
    online-platform folder), so this just leaves a record that the handoff happened."""
    from database import get_job_detail
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    content = (
        f"Sent to Resume Tailor (Dashboard) at {datetime.now().isoformat()}\n"
        f"Job: {job.get('title', '')} at {job.get('company', '')}\n"
    )
    _write_job_artifact(job_id, job.get('company', ''), "send_to_resume_tailor.txt", content)
    return {"status": "logged"}


DRAFT_TYPES = ["recruiter_email", "follow_up_email", "linkedin_message"]

class GenerateDraftRequest(BaseModel):
    draft_type: str

@app.post("/api/jobs/{job_id}/draft")
@limiter.limit("15/minute")
async def generate_job_draft_endpoint(request: Request, job_id: int, payload: GenerateDraftRequest):
    if payload.draft_type not in DRAFT_TYPES:
        raise HTTPException(status_code=400, detail=f"draft_type must be one of {DRAFT_TYPES}")
    from database import get_job_detail, save_job_draft
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    title = job.get('title') or 'this role'
    company = job.get('company') or 'the company'
    candidate_profile = _load_candidate_profile()

    try:
        if payload.draft_type == "recruiter_email":
            desc = (job.get('description') or '').strip()
            if len(desc) < 40:
                raise HTTPException(status_code=400, detail="Not enough job description to write from. Try 'Fetch job details' first.")
            draft = generate_recruiter_outreach_email(title, company, desc, candidate_profile)
        elif payload.draft_type == "follow_up_email":
            days_since = 0
            if job.get('status_updated_at'):
                try:
                    days_since = (datetime.now() - datetime.fromisoformat(job['status_updated_at'])).days
                except (ValueError, TypeError):
                    days_since = 0
            draft = generate_checkin_followup_email(title, company, days_since, candidate_profile)
        else:  # linkedin_message
            draft = generate_linkedin_message(title, company, candidate_profile)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=f"The AI provider is unavailable: {e}")

    saved = save_job_draft(job_id, payload.draft_type, draft.get('subject', ''), draft.get('body', ''))

    artifact_filenames = {
        "recruiter_email": "recruiter_email.txt",
        "follow_up_email": "follow_up_email.txt",
        "linkedin_message": "linkedin_message.txt",
    }
    artifact_content = f"Subject: {saved.get('subject', '')}\n\n{saved.get('body', '')}\n"
    _write_job_artifact(job_id, company, artifact_filenames[payload.draft_type], artifact_content)

    return {"draft_type": payload.draft_type, "draft": saved}


class SaveDraftRequest(BaseModel):
    draft_type: str
    subject: str = ""
    body: str

@app.post("/api/jobs/{job_id}/draft/save")
@limiter.limit("30/minute")
async def save_job_draft_endpoint(request: Request, job_id: int, payload: SaveDraftRequest):
    if payload.draft_type not in DRAFT_TYPES:
        raise HTTPException(status_code=400, detail=f"draft_type must be one of {DRAFT_TYPES}")
    from database import get_resume_by_id, save_job_draft
    if not get_resume_by_id(job_id):
        raise HTTPException(status_code=404, detail="Job not found")
    saved = save_job_draft(job_id, payload.draft_type, payload.subject, payload.body)
    return {"draft_type": payload.draft_type, "draft": saved}


class SaveNotesRequest(BaseModel):
    notes: str

@app.post("/api/jobs/{job_id}/notes")
@limiter.limit("30/minute")
async def save_job_notes_endpoint(request: Request, job_id: int, payload: SaveNotesRequest):
    if len(payload.notes) > 2000:
        raise HTTPException(status_code=400, detail="Notes must be 2000 characters or fewer")
    from database import get_resume_by_id, update_user_notes, get_job_detail
    if not get_resume_by_id(job_id):
        raise HTTPException(status_code=404, detail="Job not found")
    update_user_notes(job_id, payload.notes)
    return get_job_detail(job_id)


def _guess_domain_from_url(url: str) -> str:
    if not url:
        return ""
    m = re.search(r'https?://(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z0-9.-]+)', url)
    if not m:
        return ""
    domain = m.group(1).lower()
    # Job-board/ATS domains aren't the employer's own domain — skip those, we only
    # want to guess an email/careers-page domain when the link is the company's own site.
    board_domains = {'linkedin.com', 'dice.com', 'indeed.com', 'ziprecruiter.com', 'glassdoor.com',
                      'greenhouse.io', 'lever.co', 'myworkdayjobs.com', 'workday.com', 'icims.com',
                      'smartrecruiters.com', 'jobvite.com', 'monster.com'}
    if any(domain == b or domain.endswith('.' + b) for b in board_domains):
        return ""
    return domain

def _scrape_contact_candidates(company: str, title: str) -> str:
    """Free contact discovery, step 1: DDGS search + scrape (same approach the
    auto-search fallback already uses) for recruiter/talent-acquisition mentions and
    the company's own careers/contact page. Returns combined scraped text with source
    URL markers for the AI extraction step, or "" if nothing could be fetched. This
    does NOT call any AI — it's plain search + scraping, same as the rest of the app."""
    if not company:
        return ""
    queries = [
        # Careers/official-site query goes first — it's the highest-value source (gets
        # us the real company domain, which everything else — email guess, careers page
        # link — depends on) but was previously missing entirely; the other two queries
        # only ever found generic recruiter mentions, rarely the company's own site.
        f'"{company}" official website OR careers',
        f'"{company}" recruiter OR "talent acquisition"',
        f'"{company}" hiring manager {title}'.strip(),
    ]
    snippets = []
    seen_urls = set()
    for q in queries:
        try:
            results = DDGS().text(q, max_results=3)
        except Exception as e:
            logger.warning(f"Contact search failed for '{q}': {e}")
            continue
        for r in results:
            url = r.get("href", "")
            if not url or url in seen_urls:
                continue
            seen_urls.add(url)
            try:
                headers = {"User-Agent": "Mozilla/5.0"}
                resp = requests.get(url, headers=headers, timeout=5)
                if resp.status_code == 200:
                    soup = BeautifulSoup(resp.text, 'html.parser')
                    text = soup.get_text(separator=' ', strip=True)[:2500]
                    snippets.append(f"SOURCE: {url}\n{text}")
            except Exception as e:
                logger.warning(f"Contact page fetch failed for {url}: {e}")
            if len(snippets) >= 6:
                break
        if len(snippets) >= 6:
            break
    return "\n\n---\n\n".join(snippets)


def _find_company_domain(company: str, job_url: str = "", employer_website: str = "") -> str:
    """Find the employer's own domain, trying the most reliable source first:
    (1) employer_website, if JSearch/Claude already gave us one — no guessing needed.
    (2) the job posting URL itself — works when it's a company ATS page, not a job board.
    (3) a quick DDGS search for the company's official site — most postings route
    through a job board (LinkedIn/Indeed/Dice/etc.) whose domain _guess_domain_from_url
    deliberately rejects, so without this fallback the domain comes back empty for the
    large majority of postings and every field derived from it (careers page, email
    guess) goes blank too.
    Returns "" if nothing usable found — never fabricates a domain."""
    domain = _guess_domain_from_url(employer_website)
    if domain:
        return domain
    domain = _guess_domain_from_url(job_url)
    if domain:
        return domain
    if not company:
        return ""
    try:
        results = DDGS().text(f'"{company}" official website', max_results=3)
        for r in results:
            candidate = _guess_domain_from_url(r.get("href", ""))
            if candidate:
                return candidate
    except Exception as e:
        logger.warning(f"Company domain search failed for '{company}': {e}")
    return ""


@app.post("/api/jobs/{job_id}/contact")
@limiter.limit("30/minute")
async def find_job_contact_endpoint(request: Request, job_id: int):
    """Best-effort contact discovery in two cheap steps, NOT live AI web search:
    (1) free DDGS search + page scraping (same approach auto-search already uses) for
    recruiter/talent-acquisition mentions and the company's careers/contact page, then
    (2) one Gemini call to ORGANIZE/EXTRACT any real names/emails already present in
    that scraped text — never to search the web itself or invent a person. Also builds
    deterministic LinkedIn search links and a plausible email guess as a fallback.
    Every result is labeled with where it came from so nothing looks more verified
    than it is."""
    from database import get_resume_by_id, save_job_contact, get_job_detail
    job = get_job_detail(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    company = (job.get('company') or '').strip()
    title = (job.get('title') or '').strip()
    domain = _find_company_domain(company, job.get('url') or '', job.get('employer_website') or '')
    company_slug = re.sub(r'[^a-z0-9]+', '-', company.lower()).strip('-')

    import urllib.parse
    recruiter_kw = urllib.parse.quote(f"{company} recruiter")
    hiring_manager_kw = urllib.parse.quote(f"{company} hiring manager {title}".strip())

    # Preserve any real contact info the scoring pass already pulled out of the
    # posting text — this call adds LinkedIn/careers-page suggestions on top of it,
    # it shouldn't overwrite a real extracted name/email with a guessed one.
    scraped_name = (job.get('contact_name') or '').strip()
    scraped_email = (job.get('contact_email') or '').strip()
    scraped_phone = (job.get('contact_phone') or '').strip()

    # Step 1 + 2: free search/scrape, then AI-organize whatever text that turns up.
    # found_contacts entries are labeled "found_on_web" (from company/contact pages)
    # to distinguish them from scraped_name/scraped_email above, which are labeled
    # "verified_from_posting" (pulled straight from the job description itself).
    found_contacts = []
    company_page_url = None
    try:
        scraped_text = _scrape_contact_candidates(company, title)
        if scraped_text:
            extraction = extract_contacts_from_text(company, title, scraped_text)
            for c in (extraction.get('contacts') or []):
                if c.get('name') or c.get('email'):
                    found_contacts.append({**c, 'provenance': 'found_on_web'})
            company_page_url = (extraction.get('company_page_url') or '').strip() or None
    except Exception as e:
        logger.warning(f"Free contact discovery failed for job {job_id}: {e}")

    contact = {
        "contact_name": scraped_name or None,
        "contact_phone": scraped_phone or None,
        "contact_provenance": "verified_from_posting" if (scraped_name or scraped_email) else None,
        "found_contacts": found_contacts,
        "linkedin_recruiter_search": f"https://www.linkedin.com/search/results/people/?keywords={recruiter_kw}",
        "linkedin_hiring_manager_search": f"https://www.linkedin.com/search/results/people/?keywords={hiring_manager_kw}",
        "linkedin_company_url": f"https://www.linkedin.com/company/{company_slug}" if company_slug else None,
        "careers_page": company_page_url or (f"https://{domain}/careers" if domain else None),
        "email_guess": scraped_email or (found_contacts[0].get('email') if found_contacts and found_contacts[0].get('email') else None) or (f"careers@{domain}" if domain else None),
        "outreach_strategy": (
            f"Search LinkedIn for a recruiter or hiring manager at {company or 'this company'}, "
            f"send a short connection note referencing the {title or 'role'} posting, then follow up "
            f"by email if there's no response within 3-5 business days. Verify any guessed email/domain "
            f"before sending — these are suggestions, not confirmed contacts."
        ),
        "verified": bool(scraped_name or scraped_email),
    }
    saved = save_job_contact(job_id, contact)

    contact_lines = [f"Contact Info — {title or 'this role'} at {company or 'the company'}", ""]
    if saved.get('verified'):
        # verified is True when a name OR email was scraped directly from the posting
        # (see contact['verified'] above) — the old condition only checked
        # contact_name/contact_phone, so a posting with a scraped email but no name/
        # phone (e.g. "apply by emailing jane@company.com") silently lost its
        # "verified" line here even though the UI correctly showed it as verified.
        contact_lines.append(
            f"Verified from job posting: {saved.get('contact_name') or 'N/A'} — {saved.get('contact_phone') or 'N/A'}"
            + (f" — {scraped_email}" if scraped_email else "")
        )
    for c in saved.get('found_contacts') or []:
        contact_lines.append(f"Found on web: {c.get('name', '')} — {c.get('title', '')} — {c.get('email', '')} (source: {c.get('source_url', '')})")
    contact_lines.append(f"Email guess: {saved.get('email_guess') or 'N/A'}")
    contact_lines.append(f"Careers page: {saved.get('careers_page') or 'N/A'}")
    contact_lines.append(f"LinkedIn recruiter search: {saved.get('linkedin_recruiter_search') or 'N/A'}")
    contact_lines.append(f"LinkedIn hiring manager search: {saved.get('linkedin_hiring_manager_search') or 'N/A'}")
    contact_lines.append("")
    contact_lines.append(f"Outreach strategy: {saved.get('outreach_strategy') or ''}")
    _write_job_artifact(job_id, company, "contact_info.txt", "\n".join(contact_lines) + "\n")

    return {"contact": saved}


from typing import List, Optional

class AutoSearchRequest(BaseModel):
    query: str = ""
    platforms: Optional[List[str]] = ["linkedin", "dice", "indeed", "ziprecruiter"]
    work_types: Optional[List[str]] = []
    contract_types: Optional[List[str]] = []

DEVOPS_RELATED_TITLES = [
    "DevOps Engineer",
    "Site Reliability Engineer",
    "SRE",
    "Cloud Engineer",
    "Platform Engineer",
    "Infrastructure Engineer",
    "Cloud DevOps Engineer",
    "DevSecOps Engineer",
]
DEVOPS_TRIGGER_WORDS = [
    "devops", "sre", "site reliability", "platform engineer",
    "cloud engineer", "infrastructure engineer", "devsecops",
]

def _broaden_devops_query(query: str) -> str:
    """If the query is DevOps-flavored, widen it to an OR of closely related
    role titles so the search isn't limited to the literal keyword."""
    q = (query or "").strip()
    if not q:
        return q
    q_lower = q.lower()
    if not any(w in q_lower for w in DEVOPS_TRIGGER_WORDS):
        return q

    terms = [q] + [t for t in DEVOPS_RELATED_TITLES if t.lower() not in q_lower and q_lower not in t.lower()]
    seen = set()
    uniq = []
    for t in terms:
        key = t.lower()
        if key not in seen:
            seen.add(key)
            uniq.append(t)
    return "(" + " OR ".join(f'"{t}"' for t in uniq[:6]) + ")"


def _load_candidate_profile() -> str:
    profile_path = os.path.join(DATA_DIR, "profile.txt")
    if os.path.exists(profile_path):
        try:
            with open(profile_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception:
            return ""
    return ""


_JOB_TEXT_FIELD_LABELS = ['Title:', 'Company:', 'Location:', 'URL:', 'Employer Website:', 'Description:', 'Content:']


def _extract_field_from_job_text(text: str, labels) -> str:
    """Pull the value of the first matching 'Label:' line out of a job_text blob (up to
    the next recognized field label, or end of string). Used to recover the ORIGINAL
    scraped description/url/employer_website after Claude scoring — Claude is not
    asked to reproduce these fields itself (see _score_jobs_with_claude), so matching
    back to the source via posting_index is exact and has zero output-token
    truncation risk, no matter how long the original text is."""
    if isinstance(labels, str):
        labels = [labels]
    for label in labels:
        marker = f"\n{label}"
        idx = text.find(marker)
        if idx != -1:
            start = idx + len(marker)
        elif text.startswith(label):
            start = len(label)
        else:
            continue
        end = len(text)
        for kl in _JOB_TEXT_FIELD_LABELS:
            pos = text.find(f"\n{kl}", start)
            if pos != -1 and pos < end:
                end = pos
        return text[start:end].strip()
    return ""


def _score_jobs_with_claude(job_texts: list, query: str) -> list:
    """Shared Claude scoring/ranking call used by both auto-search and single-job
    re-scoring. Returns a list of job dicts (title/company/location/score/tags/reasons/
    next_action/description/url/employer_website). Does not persist anything — callers
    decide what to do with the result. Raises RuntimeError if ANTHROPIC_API_KEY is not
    set.

    Claude is NOT asked to reproduce description/url/employer_website in its own
    output — earlier versions did, which meant a long JD either burned huge amounts of
    output tokens or got silently truncated once max_tokens ran out, no matter how
    generous the *input* truncation was. Instead, each posting is tagged with a
    [POSTING #N] marker; Claude only has to echo back which N a job came from
    (posting_index), and the backend pulls description/url/employer_website straight
    from the original job_texts entry. This is exact (no LLM copy-through risk) and
    much cheaper (no giant verbatim JD in the response)."""
    import anthropic

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY is not set")

    indexed_texts = [f"[POSTING #{i}]\n{jt}" for i, jt in enumerate(job_texts)]
    context = "\n\n---\n\n".join(indexed_texts)

    candidate_profile = _load_candidate_profile()

    profile_block = (
        f"\nCandidate background (use this to judge fit, not just the query text):\n{candidate_profile}\n"
        if candidate_profile else ""
    )

    client = anthropic.Anthropic(api_key=api_key)
    prompt = f'''You are a job search assistant for a candidate whose primary target roles are DevOps / Cloud
Engineering (CI/CD, Kubernetes, Terraform/IaC, AWS/Azure/GCP, containerization, observability, automation).
I have scraped some real job postings from the web based on the query: "{query}". Each posting below is labeled
with a "[POSTING #N]" marker.
{profile_block}
Here is the scraped text from the job pages:

{context}

Extract the real job postings from this text. For each job you find, create a JSON object and set `posting_index`
to the exact N from its "[POSTING #N]" marker — this is how the app matches your result back to the original
posting's full text, so it must be accurate. Do not include description, url, or employer_website fields — the app
already has the original text and will attach those itself.
Estimate a match score (0-100) weighted toward DevOps/Cloud relevance and the candidate's background above (experience level,
tooling overlap, work authorization/relocation constraints if listed). Provide 3 brief reasons that reference specific
requirements from the posting, not generic filler.
Also provide a `next_action`: one short, concrete instruction for what the candidate should do next about this specific
posting (e.g. "Apply now — strong match", "Tailor resume first — highlight Kubernetes", "Save for later — good but not urgent").
Base it on the score and reasons, not a generic phrase.
Also look for a recruiter/hiring contact in the scraped text (a name, an email address, or a phone number near words
like "contact", "recruiter", "hiring manager", "apply by emailing", a signature block, etc.) and fill in contact_name,
contact_email, contact_phone with EXACTLY what's in the text. Leave any of these as an empty string if not literally
present — never invent or guess a contact.'''

    response = client.messages.create(
        model="claude-sonnet-5",
        # No verbatim job descriptions in the output anymore (see docstring), so this
        # comfortably covers up to 8 jobs' worth of reasons/tags/contact info.
        max_tokens=8000,
        thinking={"type": "disabled"},
        messages=[{"role": "user", "content": prompt}],
        output_config={
            "format": {
                "type": "json_schema",
                "schema": {
                    "type": "object",
                    "properties": {
                        "jobs": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "posting_index": {"type": "integer"},
                                    "title": {"type": "string"},
                                    "company": {"type": "string"},
                                    "location": {"type": "string"},
                                    "score": {"type": "integer"},
                                    "tags": {"type": "array", "items": {"type": "string"}},
                                    "reasons": {"type": "array", "items": {"type": "string"}},
                                    "next_action": {"type": "string"},
                                    "contact_name": {"type": "string"},
                                    "contact_email": {"type": "string"},
                                    "contact_phone": {"type": "string"}
                                },
                                "required": ["posting_index", "title", "company", "location", "score", "tags", "reasons", "next_action", "contact_name", "contact_email", "contact_phone"],
                                "additionalProperties": False
                            }
                        }
                    },
                    "required": ["jobs"],
                    "additionalProperties": False
                }
            }
        }
    )

    try:
        from services.usage_tracker import log_api_call
        log_api_call("claude-sonnet-5", "auto_search_scoring",
                     input_tokens=response.usage.input_tokens,
                     output_tokens=response.usage.output_tokens)
    except Exception as e:
        logger.warning(f"Failed to log claude usage: {e}")

    if response.stop_reason == "max_tokens":
        logger.warning(
            f"Claude scoring response hit max_tokens ({response.usage.output_tokens} tokens) — "
            "output was likely truncated mid-JSON; consider fewer jobs per batch."
        )

    resp_text = next(b.text for b in response.content if b.type == "text")
    data = json.loads(resp_text.strip())
    jobs = data.get('jobs', []) if isinstance(data, dict) else []

    for j in jobs:
        idx = j.pop('posting_index', None)
        raw = job_texts[idx] if isinstance(idx, int) and 0 <= idx < len(job_texts) else ""
        j['description'] = _extract_field_from_job_text(raw, ['Description:', 'Content:'])
        j['url'] = _extract_field_from_job_text(raw, 'URL:')
        j['employer_website'] = _extract_field_from_job_text(raw, 'Employer Website:')

        # Detect employment type from the FULL scraped blob (raw) — the exact same text
        # _check_c2c_c2h_only evaluated to decide whether this posting even reached
        # scoring — and persist it now, rather than recomputing later from just the
        # narrower saved `description` field. Recomputing from `description` alone can
        # under-detect: a contract-type signal ("C2C", "1099", "W2") sometimes sits
        # outside the specific "Description:" section of a scraped page (e.g. a "Job
        # Type:" sidebar field), so the badge could show "Unspecified" even though the
        # hard-reject rule that let the posting through actually saw a real signal.
        employment_type, _ = _check_employment_type(raw)
        j['employment_type'] = employment_type
        j['employment_type_label'] = EMPLOYMENT_TYPE_LABELS.get(employment_type, 'Unspecified')

    return jobs


async def _run_command_center_search(query: str, platforms: list, work_types: list, contract_types: list) -> dict:
    """Core Command Center auto-search pipeline: scrape (JSearch or DDG fallback) ->
    hard-reject rules -> Claude scoring -> persist. Shared by the HTTP endpoint and the
    daily scheduled search so neither duplicates this logic."""
    query = query if query else "DevOps Engineer"
    search_terms = _broaden_devops_query(query)
    job_texts = []
    skipped_seen_count = 0
    rapidapi_key = os.getenv("RAPIDAPI_KEY")
    jsearch_success = False
    platforms = platforms if platforms else ["linkedin", "dice", "indeed", "ziprecruiter"]
    work_types = work_types or []
    contract_types = contract_types or []

    from database import command_center_job_seen

    # Progressive recency window: try the freshest postings first, widening only when
    # needed. Scraping now walks real result pages (page=1,2,3,...) until JSearch
    # returns no results for that query/tier. Raw scraping is intentionally uncapped;
    # AI scoring remains capped later so discovery can be broad without runaway model
    # token spend.
    JSEARCH_DATE_TIERS = ["today", "3days", "week", "month"]
    TARGET_PASSING_JOBS = 8
    MAX_JOBS_TO_SCORE = 8
    MAX_JSEARCH_PAGES_PER_QUERY_TIER = 10
    MAX_STALE_PAGES_PER_QUERY_TIER = 2
    passing_job_texts = []
    rule_rejected = []


    from services.usage_tracker import get_jsearch_usage_this_month, JSEARCH_FREE_MONTHLY_LIMIT
    jsearch_quota_used = get_jsearch_usage_this_month() if rapidapi_key else 0
    if rapidapi_key and jsearch_quota_used >= JSEARCH_FREE_MONTHLY_LIMIT:
        logger.info(
            f"JSearch free-tier monthly quota ({JSEARCH_FREE_MONTHLY_LIMIT}) already reached "
            f"({jsearch_quota_used} used) — skipping JSearch and falling back to DDG scraping"
        )

    if rapidapi_key and jsearch_quota_used < JSEARCH_FREE_MONTHLY_LIMIT:
        import requests
        from services.usage_tracker import log_api_call
        url = "https://jsearch.p.rapidapi.com/search"

        modifiers = list(work_types)
        full_query = f"{search_terms} in USA"
        if modifiers:
            full_query += " " + " ".join(modifiers)

        allowed_contracts = {t.lower() for t in contract_types}
        contract_queries = []
        if allowed_contracts <= {"c2c", "c2h"} or {"c2c", "c2h"} & allowed_contracts:
            contract_terms = []
            if "c2c" in allowed_contracts or not allowed_contracts:
                contract_terms.extend(["C2C", "corp to corp", "1099 contractor"])
            if "c2h" in allowed_contracts or not allowed_contracts:
                contract_terms.extend(["C2H", "contract to hire"])
            contract_queries = [f"{search_terms} {term} remote United States" for term in contract_terms]

        # Broad query first, then explicit contract phrases. Targeted queries are what
        # usually surface C2C/C2H postings that broad job-board search buries.
        jsearch_queries = [full_query] + contract_queries

        headers = {
            "X-RapidAPI-Key": rapidapi_key,
            "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
        }
        seen_urls_this_run = set()
        jsearch_call_count = 0
        quota_hit = False
        for search_idx, active_query in enumerate(jsearch_queries):
            if quota_hit or len(passing_job_texts) >= TARGET_PASSING_JOBS:
                break
            for tier in JSEARCH_DATE_TIERS:
                if quota_hit or len(passing_job_texts) >= TARGET_PASSING_JOBS:
                    break
                page_num = 1
                seen_page_fingerprints = set()
                stale_pages = 0
                while True:
                    if len(passing_job_texts) >= TARGET_PASSING_JOBS:
                        break
                    if page_num > MAX_JSEARCH_PAGES_PER_QUERY_TIER:
                        logger.info(f"JSearch query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' stopped at page {page_num}: page safety limit")
                        break
                    if stale_pages >= MAX_STALE_PAGES_PER_QUERY_TIER:
                        logger.info(f"JSearch query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' stopped at page {page_num}: no new usable jobs on recent pages")
                        break
                    if jsearch_quota_used + jsearch_call_count >= JSEARCH_FREE_MONTHLY_LIMIT:
                        logger.info(
                            f"JSearch free-tier monthly quota ({JSEARCH_FREE_MONTHLY_LIMIT}) reached mid-run "
                            f"at query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' page {page_num} — stopping further JSearch calls"
                        )
                        quota_hit = True
                        break
                    try:
                        # Real pagination: request exactly one provider page at a time
                        # and increment `page` until the API returns an empty page.
                        querystring = {
                            "query": active_query,
                            "page": str(page_num),
                            "num_pages": "1",
                            "remote_jobs_only": "true" if search_idx == 0 else "false",
                            "date_posted": tier,
                        }
                        resp = requests.get(url, headers=headers, params=querystring, timeout=25)
                        jsearch_call_count += 1
                        log_api_call("jsearch-api", "auto_search", input_tokens=1, output_tokens=0)
                        if resp.status_code != 200:
                            logger.warning(f"JSearch API returned non-200 status for query {search_idx + 1}, tier '{tier}', page {page_num}: {resp.status_code} - {resp.text[:300]}")
                            break

                        jdata = resp.json().get("data", [])
                        if not jdata:
                            logger.info(f"JSearch query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' ended at page {page_num}: no results")
                            break

                        page_fingerprint = tuple((j.get("job_apply_link") or j.get("job_id") or j.get("job_title") or "") for j in jdata)
                        if page_fingerprint in seen_page_fingerprints:
                            logger.info(f"JSearch query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' stopped at page {page_num}: repeated provider page")
                            break
                        seen_page_fingerprints.add(page_fingerprint)

                        new_usable_this_page = 0
                        for j in jdata:
                            apply_link = j.get("job_apply_link", "")
                            if apply_link and apply_link in seen_urls_this_run:
                                continue
                            title = j.get("job_title", "")
                            company = j.get("employer_name", "")
                            if command_center_job_seen(company, title):
                                skipped_seen_count += 1
                                if apply_link:
                                    seen_urls_this_run.add(apply_link)
                                continue
                            if apply_link:
                                seen_urls_this_run.add(apply_link)

                            city = j.get("job_city") or ""
                            state = j.get("job_state") or ""
                            location = f"{city}, {state}".strip(", ")
                            desc = (j.get("job_description") or "")[:20000]
                            employer_website = j.get("employer_website") or ""
                            jt = (
                                f"Title: {title}\nCompany: {company}\nLocation: {location}\nURL: {apply_link}\n"
                                f"Employer Website: {employer_website}\nDescription: {desc}"
                            )
                            job_texts.append(jt)
                            reason = (
                                _check_position_filled(jt)
                                or _check_lead_role(jt)
                                or _check_visa_eligibility(jt)
                                or _check_foreign_language(jt)
                                or _check_remote_only(jt)
                                or _check_c2c_c2h_only(jt, contract_types)
                                or _check_experience_years(jt)
                            )
                            if reason:
                                rule_rejected.append({
                                    'company': company or _extract_company_name(jt) or 'Unknown',
                                    'title': title or _guess_title_from_job_text(jt),
                                    'reason': reason,
                                })
                            else:
                                passing_job_texts.append(jt)
                                new_usable_this_page += 1
                                if len(passing_job_texts) >= TARGET_PASSING_JOBS:
                                    break

                        if job_texts:
                            jsearch_success = True
                        stale_pages = 0 if new_usable_this_page else stale_pages + 1
                        logger.info(
                            f"JSearch query {search_idx + 1}/{len(jsearch_queries)} tier '{tier}' page {page_num}: "
                            f"{len(jdata)} results, {len(passing_job_texts)} usable, {len(rule_rejected)} rejected, "
                            f"{skipped_seen_count} already known"
                        )
                        page_num += 1
                    except Exception as e:
                        logger.warning(f"JSearch API error for query {search_idx + 1}, tier '{tier}', page {page_num}: {e}")
                        break

    if not jsearch_success:
        sites = []
        if "linkedin" in platforms: sites.append("site:linkedin.com/jobs/view/")
        if "dice" in platforms: sites.append("site:dice.com/jobs/")
        if "indeed" in platforms: sites.append("site:indeed.com/viewjob")
        if "ziprecruiter" in platforms: sites.append("site:ziprecruiter.com/jobs")
        if not sites: sites = ["site:linkedin.com/jobs/view/"]

        site_query = " OR ".join(sites)
        search_query = f"({site_query}) {search_terms} remote United States"

        # Same progressive-recency idea as the JSearch tiers above, using DDGS' own
        # timelimit filter: d=past day, w=past week, m=past month — the closest DDGS
        # equivalent to the 24h/72h/2-3-week escalation (DDGS has no 3-day bucket).
        DDG_TIME_TIERS = ["d", "w", "m"]
        seen_urls_this_run = set()
        for tier in DDG_TIME_TIERS:
            try:
                results = DDGS().text(search_query, max_results=50, timelimit=tier)
            except Exception as e:
                logger.warning(f"DDG search failed for tier '{tier}': {e}")
                continue
            for r in results:
                page_url = r.get("href", "")
                if not page_url or page_url in seen_urls_this_run:
                    continue
                seen_urls_this_run.add(page_url)
                try:
                    headers = {"User-Agent": "Mozilla/5.0"}
                    resp = requests.get(page_url, headers=headers, timeout=5)
                    if resp.status_code == 200:
                        soup = BeautifulSoup(resp.text, 'html.parser')
                        text = soup.get_text(separator=' ', strip=True)
                        # Same reasoning as the JSearch cap above — no more Claude
                        # output-token pressure to weigh this against, so raw scraped
                        # page text (nav/ads/footer included, diluting useful content)
                        # gets a generous ceiling rather than a tight one.
                        candidate_text = f"URL: {page_url}\nContent: {text[:20000]}"
                        company_guess = _extract_company_name(candidate_text) or ''
                        title_guess = _guess_title_from_job_text(candidate_text) or ''
                        if company_guess and title_guess and command_center_job_seen(company_guess, title_guess):
                            skipped_seen_count += 1
                            continue
                        job_texts.append(candidate_text)
                except Exception as e:
                    logger.warning(f"Error fetching {page_url}: {e}")

    if not job_texts:
        message = f'No postings found for "{query}" on the selected platforms. Try different platforms, work types, or a broader search term.'
        if skipped_seen_count:
            message = f'Found {skipped_seen_count} posting(s) for "{query}", but all were already in your Command Center from a previous search. Try a broader search term to find new ones.'
        result = {
            "jobs": [],
            "count": 0,
            "rejected_count": 0,
            "skipped_seen_count": skipped_seen_count,
            "query": query,
            "message": message,
            "cached": False,
            "api_spent": True,
        }
        return result

    # JSearch filters postings as it scrapes so it can stop once enough usable jobs are
    # found. The DDG fallback appends raw job_texts, so filter those here if needed.
    if not passing_job_texts and job_texts:
        for jt in job_texts:
            reason = (
                _check_position_filled(jt)
                or _check_lead_role(jt)
                or _check_visa_eligibility(jt)
                or _check_foreign_language(jt)
                or _check_remote_only(jt)
                or _check_c2c_c2h_only(jt, contract_types)
                or _check_experience_years(jt)
            )
            if reason:
                rule_rejected.append({
                    'company': _extract_field_from_job_text(jt, 'Company:') or _extract_company_name(jt) or 'Unknown',
                    'title': _extract_field_from_job_text(jt, 'Title:') or _guess_title_from_job_text(jt),
                    'reason': reason,
                })
            else:
                passing_job_texts.append(jt)

    if rule_rejected:
        from database import save_rejected_job
        for rj in rule_rejected:
            try:
                save_rejected_job(rj)
            except Exception as e:
                logger.warning(f"Failed to persist rule-rejected job for {rj.get('company')}: {e}")

    job_texts = passing_job_texts
    if not job_texts:
        result = {
            "jobs": [],
            "count": 0,
            "rejected_count": len(rule_rejected),
            "skipped_seen_count": skipped_seen_count,
            "query": query,
            "message": (
                f'All {len(rule_rejected)} posting(s) for "{query}" were filtered out by your hard rules '
                f'(visa/lead-level/experience/language/remote-only/C2C-C2H-only) — see Jobs to Skip for why. Try a broader search term.'
            ),
            "cached": False,
            "api_spent": True,
        }
        return result

    if len(job_texts) > MAX_JOBS_TO_SCORE:
        logger.info(f"Scoring first {MAX_JOBS_TO_SCORE} of {len(job_texts)} passing postings to control AI token spend")
        job_texts = job_texts[:MAX_JOBS_TO_SCORE]

    jobs = _score_jobs_with_claude(job_texts, query)

    from database import save_found_job
    for j in jobs:
        try:
            save_found_job(j)
        except Exception as e:
            logger.warning(f"Failed to persist found job for {j.get('company')}: {e}")

    if jobs:
        try:
            from services.telegram_notifier import notify_new_action_items
            notify_new_action_items()
        except Exception as e:
            logger.warning(f"Failed to send Telegram notification for new matches: {e}")

    try:
        from services.scan_status import save_last_scan
        save_last_scan(platforms, query, len(jobs))
    except Exception as e:
        logger.warning(f"Failed to save last scan status: {e}")

    message = None if jobs else f'No postings could be extracted for "{query}". Try a different search term.'
    if rule_rejected and jobs:
        message = f'{len(rule_rejected)} posting(s) filtered out by your hard rules (visa/lead-level/experience/language/remote-only/C2C-C2H-only) — see Jobs to Skip.'
    if skipped_seen_count and jobs:
        already_seen_note = f'{skipped_seen_count} already-seen posting(s) skipped (no re-scan).'
        message = f'{message} {already_seen_note}' if message else already_seen_note

    return {
        "jobs": jobs,
        "count": len(jobs),
        "rejected_count": len(rule_rejected),
        "skipped_seen_count": skipped_seen_count,
        "query": query,
        "message": message,
        "cached": False,
        "api_spent": True,
    }


@app.post("/api/jobs/auto-search")
@limiter.limit("5/minute")
async def auto_search_jobs(request: Request, payload: AutoSearchRequest):
    try:
        return await _run_command_center_search(payload.query, payload.platforms, payload.work_types, payload.contract_types)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Auto search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class ApplicationSaveRequest(BaseModel):
    company: str
    source: str = "command-center"
    status: str = "Shortlisted"
    notes: str = ""
    title: str = ""
    url: str = ""

@app.post("/api/applications")
@limiter.limit("30/minute")
async def save_application(request: Request, payload: ApplicationSaveRequest):
    if not payload.company or not payload.company.strip():
        raise HTTPException(status_code=400, detail="Company name is required")
    from database import save_application_from_job
    record_id = save_application_from_job(
        company_name=payload.company.strip(),
        status=payload.status or "Shortlisted",
        source=payload.source or "command-center",
        notes=payload.notes or "",
        title=payload.title or "",
        url=payload.url or "",
    )
    content = (
        f"Saved to Applications at {datetime.now().isoformat()}\n"
        f"Job: {payload.title or ''} at {payload.company.strip()}\n"
        f"URL: {payload.url or 'N/A'}\n"
    )
    _write_job_artifact(record_id, payload.company.strip(), "save_to_applications.txt", content)
    return {"id": record_id, "status": "saved"}


if __name__ == "__main__":
    # Don't use reload in production
    reload = "--reload" in sys.argv or os.getenv("ENVIRONMENT") == "development"
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=reload)

