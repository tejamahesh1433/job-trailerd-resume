import docx
from io import BytesIO
import difflib
import copy
from docx.text.paragraph import Paragraph

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(BytesIO(file_bytes))
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())

    return '\n'.join(full_text)

def _normalize(text: str) -> str:
    """Normalize whitespace for comparison."""
    return ' '.join(text.split())

def _find_best_paragraph_match(paragraphs, original_text):
    """Find the paragraph that best matches the original text using fuzzy matching."""
    norm_original = _normalize(original_text)
    best_para = None
    best_ratio = 0.0

    for para in paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue

        # Exact match
        if original_text in para_text:
            return para, original_text

        # Normalized exact match
        norm_para = _normalize(para_text)
        if norm_original in norm_para:
            return para, original_text

        # Fuzzy match — the paragraph text should contain something very similar
        ratio = difflib.SequenceMatcher(None, norm_original, norm_para).ratio()
        if ratio > best_ratio and ratio >= 0.75:
            best_ratio = ratio
            best_para = para

        # Also check if original is a substantial substring match
        if len(norm_original) > 20:
            # Try matching the first 60 chars as anchor
            anchor = norm_original[:60]
            if anchor in norm_para:
                return para, original_text

    return best_para, original_text

def replace_text_in_paragraph(paragraph, original_text, new_text):
    """Exact/normalized match ONLY — no fuzzy fallback here. create_tailored_docx's
    first pass calls this expecting purely exact matching, with fuzzy matching handled
    separately (and correctly) by _find_best_paragraph_match in its second pass, which
    picks the single BEST match across the whole document. This function fuzzy-matching
    on its own used to defeat that design: the first pass iterates paragraphs in
    document order and stops at whichever paragraph FIRST crossed the 0.75 similarity
    threshold, not necessarily the correct/best one — on a resume with two similar
    bullets (e.g. "Managed AWS infrastructure..." appearing under two different jobs),
    this could silently overwrite the wrong bullet instead of the one the AI actually
    intended to replace."""
    if not paragraph.text.strip() or not original_text.strip():
        return False

    para_text = paragraph.text
    norm_para = _normalize(para_text)
    norm_original = _normalize(original_text)

    # Check if original text exists (exact or normalized)
    exact_match = original_text in para_text
    normalized_match = norm_original in norm_para

    if not exact_match and not normalized_match:
        return False

    if exact_match:
        # 1. Simple case: original_text exactly matches one run
        for run in paragraph.runs:
            if original_text in run.text:
                run.text = run.text.replace(original_text, new_text)
                return True

        # 2. Complex case: text spans multiple runs
        return _replace_across_runs(paragraph, original_text, new_text)

    if normalized_match:
        # Whitespace differs — replace the entire paragraph content preserving first run's formatting
        _replace_entire_paragraph(paragraph, new_text)
        return True

    return False

def _replace_entire_paragraph(paragraph, new_text):
    """Replace entire paragraph text while preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    # Keep first run's formatting, put all new text there
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""

def _replace_across_runs(paragraph, original_text, new_text):
    """Handle replacement when text spans multiple runs."""
    text = paragraph.text
    start_idx = text.find(original_text)
    if start_idx < 0:
        return False

    end_idx = start_idx + len(original_text)

    curr_idx = 0
    match_runs = []

    for i, run in enumerate(paragraph.runs):
        run_len = len(run.text)
        if curr_idx + run_len > start_idx and curr_idx < end_idx:
            match_runs.append((i, curr_idx))
        curr_idx += run_len

    if not match_runs:
        return False

    first_run_idx, first_run_start = match_runs[0]
    first_run = paragraph.runs[first_run_idx]
    prefix = first_run.text[:max(0, start_idx - first_run_start)]

    last_run_idx, last_run_start = match_runs[-1]
    last_run = paragraph.runs[last_run_idx]
    suffix = last_run.text[max(0, end_idx - last_run_start):]

    # Inject new text into first run
    first_run.text = prefix + new_text + (suffix if first_run_idx == last_run_idx else "")

    # Clear subsequent runs involved in the match
    for idx, _ in match_runs[1:]:
        if idx == last_run_idx:
            paragraph.runs[idx].text = suffix
        else:
            paragraph.runs[idx].text = ""

    return True

def _all_paragraphs(doc):
    """Collect all paragraphs from the document body and any tables."""
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs

def create_tailored_docx(original_bytes: bytes, replacements: list) -> BytesIO:
    doc = docx.Document(BytesIO(original_bytes))
    applied_count = 0

    all_paragraphs = _all_paragraphs(doc)

    for rep in replacements:
        orig = rep.get('original', '')
        new_txt = rep.get('new', '')
        if not orig or not new_txt:
            continue

        replaced = False

        # First pass: try exact match in all paragraphs
        for para in all_paragraphs:
            if replace_text_in_paragraph(para, orig, new_txt):
                replaced = True
                applied_count += 1
                break

        # Second pass: fuzzy match if exact failed
        if not replaced:
            best_para, _ = _find_best_paragraph_match(all_paragraphs, orig)
            if best_para:
                _replace_entire_paragraph(best_para, new_txt)
                applied_count += 1

    out_stream = BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

def insert_bullets_after(original_bytes: bytes, insertions: list) -> BytesIO:
    """Insert new bullet paragraphs into the document, each cloned right after
    its matched anchor paragraph so it inherits the same list/formatting style."""
    doc = docx.Document(BytesIO(original_bytes))
    applied_count = 0

    for ins in insertions:
        anchor_text = ins.get('anchor', '')
        new_bullet = ins.get('new_bullet', '')
        if not anchor_text or not new_bullet:
            continue

        anchor_para, _ = _find_best_paragraph_match(_all_paragraphs(doc), anchor_text)
        if anchor_para is None:
            continue

        new_p_elem = copy.deepcopy(anchor_para._p)
        anchor_para._p.addnext(new_p_elem)
        new_para = Paragraph(new_p_elem, anchor_para._parent)

        if new_para.runs:
            new_para.runs[0].text = new_bullet
            for run in new_para.runs[1:]:
                run.text = ""
        else:
            new_para.text = new_bullet

        applied_count += 1

    out_stream = BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream
